from pathlib import Path
import re
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


INPUT = Path(r"C:\Users\ACER\Downloads\LVTN_v2.docx")
OUTPUT = Path(r"C:\Users\ACER\Downloads\LVTN_v2_ThemSoDoTuanTu.docx")
IMG_DIR = Path(r"C:\Users\ACER\Downloads\banhangcodex\docs\sequence_diagrams")
FONT = "Times New Roman"


def font(name="arial.ttf", size=30):
    candidates = [
        Path(r"C:\Windows\Fonts") / name,
        Path(r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\times.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


FONT_TITLE = font("arialbd.ttf", 34)
FONT_BOX = font("arialbd.ttf", 29)
FONT_TEXT = font("arial.ttf", 28)
FONT_SMALL = font("arial.ttf", 24)
FONT_NOTE = font("arial.ttf", 23)


def set_run_font(run, size=13, bold=False, italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def wrap_lines(text, max_chars):
    lines = []
    for part in str(text).split("\n"):
        lines.extend(textwrap.wrap(part, width=max_chars) or [""])
    return lines


def text_size(draw, text, fnt):
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def draw_centered_text(draw, center_x, y, text, fnt, fill=(0, 0, 0), max_chars=22, line_gap=7):
    lines = wrap_lines(text, max_chars)
    for idx, line in enumerate(lines):
        w, h = text_size(draw, line, fnt)
        draw.text((center_x - w / 2, y + idx * (h + line_gap)), line, font=fnt, fill=fill)
    return y + len(lines) * 34


def draw_arrow(draw, x1, y, x2, label, step, dashed=False):
    label_text = f"{step}: {label}"
    if dashed:
        dash = 18
        gap = 10
        x = x1
        direction = 1 if x2 >= x1 else -1
        while (direction == 1 and x < x2) or (direction == -1 and x > x2):
            nx = x + direction * min(dash, abs(x2 - x))
            draw.line((x, y, nx, y), fill=(0, 0, 0), width=3)
            x = nx + direction * gap
    else:
        draw.line((x1, y, x2, y), fill=(0, 0, 0), width=3)

    direction = 1 if x2 > x1 else -1
    head = [(x2, y), (x2 - direction * 17, y - 9), (x2 - direction * 17, y + 9)]
    draw.polygon(head, fill=(0, 0, 0))

    lines = wrap_lines(label_text, 34)
    max_w = max(text_size(draw, line, FONT_SMALL)[0] for line in lines)
    bg_h = len(lines) * 31 + 6
    mid = (x1 + x2) / 2
    draw.rectangle((mid - max_w / 2 - 9, y - bg_h - 6, mid + max_w / 2 + 9, y - 6), fill=(255, 255, 255))
    for i, line in enumerate(lines):
        w, _ = text_size(draw, line, FONT_SMALL)
        draw.text((mid - w / 2, y - bg_h + i * 30 - 2), line, font=FONT_SMALL, fill=(0, 0, 0))


def draw_self_call(draw, x, y, label, step):
    label_text = f"{step}: {label}"
    right = x + 150
    draw.line((x, y, right, y), fill=(0, 0, 0), width=3)
    draw.line((right, y, right, y + 36), fill=(0, 0, 0), width=3)
    draw.line((right, y + 36, x + 8, y + 36), fill=(0, 0, 0), width=3)
    draw.polygon([(x, y + 36), (x + 16, y + 27), (x + 16, y + 45)], fill=(0, 0, 0))
    lines = wrap_lines(label_text, 27)
    for i, line in enumerate(lines):
        draw.text((x + 15, y - 58 + i * 28), line, font=FONT_SMALL, fill=(0, 0, 0))


def draw_actor(draw, x, y, label):
    draw.ellipse((x - 18, y, x + 18, y + 36), outline=(0, 0, 0), width=3)
    draw.line((x, y + 36, x, y + 95), fill=(0, 0, 0), width=3)
    draw.line((x - 38, y + 58, x + 38, y + 58), fill=(0, 0, 0), width=3)
    draw.line((x, y + 95, x - 36, y + 145), fill=(0, 0, 0), width=3)
    draw.line((x, y + 95, x + 36, y + 145), fill=(0, 0, 0), width=3)
    draw_centered_text(draw, x, y + 152, label, FONT_BOX, max_chars=18)


def draw_sequence(path, title, actor_label, messages, errors):
    columns = {"Actor": 210, "View": 720, "Controller": 1210, "Data": 1660}
    top = 62
    box_y = 90
    line_top = 255
    step_gap = 92
    note_height = 92 if errors else 30
    height = max(860, line_top + len(messages) * step_gap + note_height + 70)
    width = 1900

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    draw.text((45, 25), title, font=FONT_TITLE, fill=(0, 0, 0))

    # Participant boxes
    for name in ("View", "Controller", "Data"):
        x = columns[name]
        draw.rectangle((x - 105, box_y, x + 105, box_y + 66), outline=(0, 0, 0), width=3, fill=(255, 255, 255))
        draw_centered_text(draw, x, box_y + 17, name, FONT_BOX, max_chars=16)

    draw_actor(draw, columns["Actor"], box_y - 5, actor_label)

    # Lifelines
    line_bottom = height - 95
    for x in columns.values():
        y = line_top
        while y < line_bottom:
            draw.line((x, y, x, min(y + 18, line_bottom)), fill=(150, 150, 150), width=2)
            y += 33

    # Activation bars
    for name in ("View", "Controller", "Data"):
        x = columns[name]
        draw.rectangle((x - 12, line_top + 28, x + 12, line_bottom - 22), outline=(0, 0, 0), width=2, fill=(255, 255, 255))

    for idx, (src, dst, label, kind) in enumerate(messages, start=1):
        y = line_top + idx * step_gap
        if src == dst:
            draw_self_call(draw, columns[src], y, label, idx)
        else:
            draw_arrow(draw, columns[src], y, columns[dst], label, idx, dashed=(kind == "return"))

    if errors:
        note_top = line_bottom + 18
        draw.rectangle((120, note_top, width - 120, note_top + 58), outline=(80, 80, 80), width=2, fill=(246, 246, 246))
        note = "Luồng ngoại lệ: " + "; ".join(errors)
        draw.text((145, note_top + 16), note, font=FONT_NOTE, fill=(0, 0, 0))

    image.save(path, "PNG")


def m(src, dst, text, kind="call"):
    return (src, dst, text, kind)


DIAGRAMS = [
    {
        "name": "đăng ký",
        "actor": "Khách hàng",
        "messages": [
            m("Actor", "View", "Nhập thông tin đăng ký"),
            m("View", "Controller", "Gửi thông tin đăng ký"),
            m("Controller", "Controller", "Kiểm tra dữ liệu và định dạng"),
            m("Controller", "Data", "Kiểm tra email/số điện thoại"),
            m("Data", "Controller", "Trả kết quả kiểm tra", "return"),
            m("Controller", "Controller", "Mã hóa mật khẩu"),
            m("Controller", "Data", "Lưu tài khoản mới"),
            m("Data", "Controller", "Lưu thành công", "return"),
            m("Controller", "View", "Trả kết quả đăng ký", "return"),
            m("View", "Actor", "Hiển thị đăng ký thành công", "return"),
        ],
        "errors": ["Dữ liệu không hợp lệ", "Email/số điện thoại đã tồn tại"],
    },
    {
        "name": "đăng nhập",
        "actor": "Khách hàng/Nhân viên/Quản trị viên",
        "messages": [
            m("Actor", "View", "Nhập thông tin đăng nhập"),
            m("Actor", "View", "Bấm nút Đăng nhập"),
            m("View", "Controller", "Gửi yêu cầu đăng nhập"),
            m("Controller", "Controller", "Kiểm tra dữ liệu rỗng"),
            m("Controller", "Controller", "Kiểm tra định dạng"),
            m("Controller", "Data", "Kiểm tra tài khoản trong CSDL"),
            m("Data", "Controller", "Trả tài khoản", "return"),
            m("Controller", "Controller", "Kiểm tra mật khẩu"),
            m("Controller", "Controller", "Tạo token/phiên đăng nhập"),
            m("Controller", "View", "Trả kết quả thành công", "return"),
            m("View", "Actor", "Điều hướng theo vai trò", "return"),
        ],
        "errors": ["Sai tài khoản/mật khẩu", "Tài khoản không tồn tại"],
    },
    {
        "name": "chỉnh sửa thông tin cá nhân",
        "actor": "Khách hàng",
        "messages": [
            m("Actor", "View", "Mở trang thông tin cá nhân"),
            m("View", "Controller", "Yêu cầu dữ liệu hồ sơ"),
            m("Controller", "Data", "Lấy thông tin người dùng"),
            m("Data", "Controller", "Trả thông tin hồ sơ", "return"),
            m("Controller", "View", "Trả dữ liệu hiển thị", "return"),
            m("Actor", "View", "Chỉnh sửa và bấm Lưu"),
            m("View", "Controller", "Gửi dữ liệu cập nhật"),
            m("Controller", "Controller", "Kiểm tra dữ liệu"),
            m("Controller", "Data", "Cập nhật hồ sơ"),
            m("Data", "Controller", "Trả kết quả", "return"),
            m("Controller", "View", "Thông báo cập nhật thành công", "return"),
        ],
        "errors": ["Dữ liệu không hợp lệ", "Email/số điện thoại bị trùng"],
    },
    {
        "name": "tìm kiếm sản phẩm",
        "actor": "Khách hàng",
        "messages": [
            m("Actor", "View", "Nhập từ khóa/chọn bộ lọc"),
            m("View", "Controller", "Gửi điều kiện tìm kiếm"),
            m("Controller", "Controller", "Xử lý điều kiện tìm kiếm"),
            m("Controller", "Data", "Truy vấn sản phẩm phù hợp"),
            m("Data", "Controller", "Trả danh sách sản phẩm", "return"),
            m("Controller", "View", "Trả dữ liệu sản phẩm", "return"),
            m("View", "Actor", "Hiển thị kết quả tìm kiếm", "return"),
        ],
        "errors": ["Không có sản phẩm phù hợp"],
    },
    {
        "name": "quản lý giỏ hàng",
        "actor": "Khách hàng",
        "messages": [
            m("Actor", "View", "Chọn sản phẩm/biến thể"),
            m("Actor", "View", "Thêm/cập nhật/xóa sản phẩm"),
            m("View", "Controller", "Gửi sản phẩm, biến thể, số lượng"),
            m("Controller", "Data", "Kiểm tra sản phẩm tồn tại"),
            m("Data", "Controller", "Trả thông tin sản phẩm", "return"),
            m("Controller", "Data", "Kiểm tra tồn kho"),
            m("Data", "Controller", "Trả số lượng tồn", "return"),
            m("Controller", "Data", "Lưu thay đổi giỏ hàng"),
            m("Data", "Controller", "Trả kết quả", "return"),
            m("Controller", "View", "Trả giỏ hàng đã cập nhật", "return"),
            m("View", "Actor", "Hiển thị giỏ hàng", "return"),
        ],
        "errors": ["Sản phẩm không tồn tại", "Tồn kho không đủ"],
    },
    {
        "name": "thanh toán đơn hàng",
        "actor": "Khách hàng",
        "messages": [
            m("Actor", "View", "Mở trang thanh toán"),
            m("View", "Controller", "Gửi địa chỉ/phương thức thanh toán"),
            m("Controller", "Data", "Lấy cấu hình phí vận chuyển"),
            m("Data", "Controller", "Trả phí theo khu vực", "return"),
            m("Controller", "View", "Trả tổng thanh toán", "return"),
            m("Actor", "View", "Xác nhận đặt hàng"),
            m("View", "Controller", "Gửi yêu cầu tạo đơn"),
            m("Controller", "Controller", "Kiểm tra giỏ hàng và tồn kho"),
            m("Controller", "Data", "Lưu đơn hàng và chi tiết"),
            m("Data", "Controller", "Trả mã đơn hàng", "return"),
            m("Controller", "View", "Trả kết quả/QR nếu chuyển khoản", "return"),
            m("View", "Actor", "Hiển thị đặt hàng thành công", "return"),
        ],
        "errors": ["Giỏ hàng rỗng", "Tồn kho không đủ", "Địa chỉ thiếu"],
    },
    {
        "name": "yêu cầu trả hàng / hoàn tiền",
        "actor": "Khách hàng",
        "messages": [
            m("Actor", "View", "Mở chi tiết đơn đã hoàn thành"),
            m("View", "Controller", "Yêu cầu thông tin đơn hàng"),
            m("Controller", "Data", "Lấy đơn hàng và chi tiết"),
            m("Data", "Controller", "Trả dữ liệu đơn", "return"),
            m("Controller", "View", "Hiển thị nút trả hàng", "return"),
            m("Actor", "View", "Chọn sản phẩm, số lượng, lý do"),
            m("View", "Controller", "Gửi yêu cầu trả hàng"),
            m("Controller", "Controller", "Kiểm tra chủ đơn, trạng thái, số lượng"),
            m("Controller", "Data", "Lưu yêu cầu trả hàng"),
            m("Data", "Controller", "Trả kết quả", "return"),
            m("Controller", "View", "Thông báo gửi yêu cầu thành công", "return"),
        ],
        "errors": ["Đơn chưa hoàn thành", "Đơn không thuộc khách hàng", "Số lượng trả không hợp lệ"],
    },
    {
        "name": "xử lý yêu cầu trả hàng",
        "actor": "Nhân viên/Quản trị viên",
        "messages": [
            m("Actor", "View", "Mở danh sách yêu cầu trả hàng"),
            m("View", "Controller", "Yêu cầu danh sách"),
            m("Controller", "Data", "Lấy danh sách yêu cầu"),
            m("Data", "Controller", "Trả danh sách", "return"),
            m("Controller", "View", "Hiển thị danh sách", "return"),
            m("Actor", "View", "Xem chi tiết và duyệt/từ chối"),
            m("View", "Controller", "Gửi quyết định xử lý"),
            m("Controller", "Controller", "Kiểm tra quyền và trạng thái"),
            m("Controller", "Data", "Cập nhật trạng thái/yêu cầu"),
            m("Controller", "Data", "Cập nhật tồn kho nếu hoàn tất"),
            m("Data", "Controller", "Lưu thay đổi", "return"),
            m("Controller", "View", "Hiển thị kết quả xử lý", "return"),
        ],
        "errors": ["Không đủ quyền", "Yêu cầu đã được xử lý"],
    },
    {
        "name": "quản lý danh mục",
        "actor": "Nhân viên/Quản trị viên",
        "messages": [
            m("Actor", "View", "Mở trang quản lý danh mục"),
            m("View", "Controller", "Yêu cầu danh sách danh mục"),
            m("Controller", "Data", "Lấy danh sách danh mục"),
            m("Data", "Controller", "Trả danh sách", "return"),
            m("Controller", "View", "Hiển thị danh mục", "return"),
            m("Actor", "View", "Thêm/sửa/ẩn danh mục"),
            m("View", "Controller", "Gửi dữ liệu danh mục"),
            m("Controller", "Controller", "Kiểm tra dữ liệu và trùng tên"),
            m("Controller", "Data", "Lưu thay đổi danh mục"),
            m("Data", "Controller", "Trả kết quả", "return"),
            m("Controller", "View", "Thông báo thành công", "return"),
        ],
        "errors": ["Tên danh mục rỗng", "Danh mục đang được sản phẩm sử dụng"],
    },
    {
        "name": "quản lý sản phẩm",
        "actor": "Nhân viên/Quản trị viên",
        "messages": [
            m("Actor", "View", "Mở trang quản lý sản phẩm"),
            m("View", "Controller", "Yêu cầu danh sách sản phẩm"),
            m("Controller", "Data", "Truy vấn sản phẩm"),
            m("Data", "Controller", "Trả danh sách", "return"),
            m("Controller", "View", "Hiển thị danh sách", "return"),
            m("Actor", "View", "Thêm/sửa sản phẩm"),
            m("View", "Controller", "Gửi thông tin sản phẩm và biến thể"),
            m("Controller", "Controller", "Kiểm tra dữ liệu, giá, danh mục"),
            m("Controller", "Data", "Lưu sản phẩm"),
            m("Controller", "Data", "Lưu biến thể/thuộc tính"),
            m("Data", "Controller", "Trả kết quả", "return"),
            m("Controller", "View", "Thông báo thành công", "return"),
        ],
        "errors": ["Thiếu tên sản phẩm", "Giá không hợp lệ", "Danh mục không tồn tại"],
    },
    {
        "name": "quản lý đơn hàng",
        "actor": "Nhân viên/Quản trị viên",
        "messages": [
            m("Actor", "View", "Mở trang quản lý đơn hàng"),
            m("View", "Controller", "Yêu cầu danh sách đơn hàng"),
            m("Controller", "Data", "Lấy danh sách đơn hàng"),
            m("Data", "Controller", "Trả danh sách", "return"),
            m("Controller", "View", "Hiển thị danh sách", "return"),
            m("Actor", "View", "Xem chi tiết/cập nhật trạng thái"),
            m("View", "Controller", "Gửi yêu cầu cập nhật"),
            m("Controller", "Data", "Lấy đơn hàng và chi tiết"),
            m("Data", "Controller", "Trả dữ liệu đơn", "return"),
            m("Controller", "Controller", "Kiểm tra quyền và luồng trạng thái"),
            m("Controller", "Data", "Lưu trạng thái mới"),
            m("Controller", "View", "Thông báo cập nhật thành công", "return"),
        ],
        "errors": ["Chuyển trạng thái sai quy trình", "Đơn đã hủy/hoàn tất"],
    },
    {
        "name": "quản lý khuyến mãi",
        "actor": "Quản trị viên",
        "messages": [
            m("Actor", "View", "Mở trang quản lý khuyến mãi"),
            m("View", "Controller", "Yêu cầu danh sách khuyến mãi"),
            m("Controller", "Data", "Lấy danh sách mã"),
            m("Data", "Controller", "Trả danh sách", "return"),
            m("Controller", "View", "Hiển thị danh sách", "return"),
            m("Actor", "View", "Thêm/sửa mã khuyến mãi"),
            m("View", "Controller", "Gửi dữ liệu khuyến mãi"),
            m("Controller", "Controller", "Kiểm tra mã, giá trị, thời gian"),
            m("Controller", "Data", "Lưu mã khuyến mãi"),
            m("Data", "Controller", "Trả kết quả", "return"),
            m("Controller", "View", "Thông báo thành công", "return"),
        ],
        "errors": ["Mã bị trùng", "Thời gian/giá trị giảm không hợp lệ"],
    },
    {
        "name": "quản lý tồn kho",
        "actor": "Nhân viên/Quản trị viên",
        "messages": [
            m("Actor", "View", "Mở trang quản lý tồn kho"),
            m("View", "Controller", "Yêu cầu danh sách tồn kho"),
            m("Controller", "Data", "Truy vấn tồn kho"),
            m("Data", "Controller", "Trả danh sách tồn", "return"),
            m("Controller", "View", "Hiển thị tồn kho", "return"),
            m("Actor", "View", "Nhập kho/điều chỉnh kho"),
            m("View", "Controller", "Gửi sản phẩm, số lượng, ghi chú"),
            m("Controller", "Controller", "Kiểm tra dữ liệu và sản phẩm"),
            m("Controller", "Data", "Cập nhật số lượng tồn"),
            m("Controller", "Data", "Ghi lịch sử biến động kho"),
            m("Data", "Controller", "Lưu thay đổi", "return"),
            m("Controller", "View", "Thông báo thành công", "return"),
        ],
        "errors": ["Số lượng không hợp lệ", "Sản phẩm không tồn tại", "Thiếu lý do điều chỉnh"],
    },
    {
        "name": "báo cáo thống kê",
        "actor": "Quản trị viên",
        "messages": [
            m("Actor", "View", "Mở trang báo cáo/thống kê"),
            m("View", "Controller", "Gửi khoảng thời gian và bộ lọc"),
            m("Controller", "Controller", "Kiểm tra điều kiện lọc"),
            m("Controller", "Data", "Truy vấn đơn hàng, doanh thu, sản phẩm"),
            m("Data", "Controller", "Trả dữ liệu thống kê", "return"),
            m("Controller", "Controller", "Tổng hợp số liệu"),
            m("Controller", "View", "Trả kết quả thống kê", "return"),
            m("View", "Actor", "Hiển thị biểu đồ và chỉ số", "return"),
        ],
        "errors": ["Khoảng thời gian không hợp lệ", "Không có dữ liệu"],
    },
]


def insert_paragraph_after(anchor, text="", style=None, alignment=None):
    paragraph = anchor._parent.add_paragraph()
    if style is not None:
        paragraph.style = style
    if text:
        run = paragraph.add_run(text)
        set_run_font(run)
    if alignment is not None:
        paragraph.alignment = alignment
    anchor._p.addnext(paragraph._p)
    return paragraph


def insert_picture_after(anchor, image_path):
    paragraph = anchor._parent.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.35))
    anchor._p.addnext(paragraph._p)
    return paragraph


def main():
    if not INPUT.exists():
        raise FileNotFoundError(INPUT)

    IMG_DIR.mkdir(parents=True, exist_ok=True)

    image_paths = []
    for idx, item in enumerate(DIAGRAMS, start=1):
        image_path = IMG_DIR / f"sequence_{idx:02d}.png"
        title = f"Sơ đồ tuần tự {item['name']}"
        draw_sequence(image_path, title, item["actor"], item["messages"], item["errors"])
        image_paths.append(image_path)

    doc = Document(INPUT)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    figure_nums = [int(match.group(1)) for match in re.finditer(r"Hình\s+3\.(\d+)", all_text)]
    next_figure = max(figure_nums or [14]) + 1

    heading = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Sơ đồ tuần tự":
            heading = paragraph
            break
    if heading is None:
        raise RuntimeError("Không tìm thấy mục Sơ đồ tuần tự.")

    anchor = heading
    intro = (
        "Các sơ đồ tuần tự dưới đây mô tả luồng xử lý theo mô hình MVC, gồm Actor, View, "
        "Controller và Data. Mỗi sơ đồ được xây dựng tương ứng với một use case chi tiết ở mục 3.2.1."
    )
    anchor = insert_paragraph_after(anchor, intro, style=heading.style)

    for i, item in enumerate(DIAGRAMS):
        # Add a small spacer to keep each diagram visually separate.
        anchor = insert_paragraph_after(anchor, "", style=heading.style)
        anchor = insert_picture_after(anchor, image_paths[i])
        caption_text = f"Hình 3.{next_figure + i}. Sơ đồ tuần tự {item['name']}"
        caption = insert_paragraph_after(anchor, caption_text, style=heading.style, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        for run in caption.runs:
            set_run_font(run, size=12, italic=True)
        anchor = caption

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
