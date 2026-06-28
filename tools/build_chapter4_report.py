from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path("Chuong_4_Thu_Nghiem_TienProSport.docx")


def set_cell_text(cell, text: str, size: float = 8.5, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_cell_margins(table, top=80, start=90, bottom=80, end=90):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], font_size: float = 8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, "D9EAF7")
        set_cell_text(cell, header, size=font_size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, len(row) - 1) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, size=font_size, align=align)

    doc.add_paragraph()
    return table


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14 if level == 1 else 13)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc: Document, text: str, bold_label: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(13)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(13)
    return p


def configure_document(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in (("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 13)):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)


scenario_rows = [
    ["TC_USER_01", "Đăng ký tài khoản", "Kiểm tra tạo tài khoản khách hàng mới.", "Tên, email, mật khẩu, số điện thoại hợp lệ.", "Mở màn đăng ký; nhập dữ liệu; xác nhận tạo tài khoản.", "Tài khoản được tạo, có giỏ hàng mặc định, có thể dùng để đăng nhập.", "Cao"],
    ["TC_USER_02", "Đăng nhập và đăng xuất", "Kiểm tra xác thực người dùng và kết thúc phiên.", "Email/mật khẩu hợp lệ.", "Đăng nhập; truy cập trang cá nhân; đăng xuất.", "Hệ thống cấp token đăng nhập, hiển thị đúng thông tin người dùng và xóa phiên khi đăng xuất.", "Cao"],
    ["TC_PRODUCT_01", "Xem, tìm kiếm, lọc sản phẩm", "Kiểm tra danh sách sản phẩm, danh mục, bộ lọc và sắp xếp.", "Từ khóa, danh mục, thương hiệu, tiêu chí sắp xếp.", "Mở trang sản phẩm; tìm kiếm; chọn danh mục; lọc thương hiệu; thay đổi sắp xếp.", "Danh sách trả về đúng sản phẩm, đúng danh mục và điều kiện lọc.", "Cao"],
    ["TC_PRODUCT_02", "Xem chi tiết sản phẩm", "Kiểm tra thông tin sản phẩm, ảnh, biến thể và tồn kho.", "Mã sản phẩm đang hoạt động.", "Mở trang chi tiết; xem ảnh; chọn biến thể; xem giá và tồn kho.", "Trang chi tiết hiển thị đủ thông tin, biến thể, thuộc tính và ảnh chính.", "Cao"],
    ["TC_CART_01", "Thêm sản phẩm vào giỏ hàng", "Kiểm tra thêm biến thể còn hàng vào giỏ.", "Mã biến thể còn tồn, số lượng hợp lệ.", "Chọn biến thể; nhập số lượng; bấm thêm vào giỏ.", "Giỏ hàng có dòng sản phẩm đúng biến thể, số lượng và thành tiền.", "Cao"],
    ["TC_CART_02", "Cập nhật và xóa giỏ hàng", "Kiểm tra thay đổi số lượng và xóa sản phẩm khỏi giỏ.", "Sản phẩm đã có trong giỏ.", "Mở giỏ hàng; tăng/giảm số lượng; xóa từng dòng.", "Số lượng, tạm tính và tổng tiền cập nhật đúng; dòng bị xóa không còn trong giỏ.", "Cao"],
    ["TC_ORDER_01", "Đặt hàng COD", "Kiểm tra tạo đơn hàng thanh toán khi nhận hàng.", "Giỏ hàng có sản phẩm, địa chỉ hợp lệ, phương thức COD.", "Mở thanh toán; nhập thông tin nhận hàng; chọn COD; đặt hàng.", "Đơn được tạo trạng thái chờ xử lý, tồn kho giảm, giỏ hàng được làm sạch.", "Cao"],
    ["TC_ORDER_02", "Đặt hàng QR chuyển khoản", "Kiểm tra luồng chuyển khoản QR và xác nhận thủ công.", "Giỏ hàng có sản phẩm, địa chỉ hợp lệ, phương thức QR.", "Đặt hàng QR; xem mã QR; bấm đã chuyển khoản; admin xác nhận.", "Đơn có nội dung chuyển khoản, link VietQR, trạng thái thanh toán chuyển sang đã thanh toán sau xác nhận.", "Cao"],
    ["TC_ORDER_03", "Lịch sử và chi tiết đơn hàng", "Kiểm tra khách chỉ xem được đơn của mình.", "Tài khoản có đơn hàng.", "Mở lịch sử đơn; mở chi tiết đơn; thử truy cập đơn của tài khoản khác.", "Danh sách/chi tiết đơn hiển thị đúng; đơn không thuộc tài khoản hiện tại bị từ chối.", "Cao"],
    ["TC_REVIEW_01", "Đánh giá sản phẩm", "Kiểm tra đánh giá sau khi đơn đã giao.", "Đơn trạng thái delivered, sản phẩm đã mua.", "Gửi đánh giá; admin duyệt; xem đánh giá công khai.", "Đánh giá chỉ hiển thị công khai sau khi admin duyệt.", "Trung bình"],
    ["TC_RETURN_01", "Gửi yêu cầu trả hàng", "Kiểm tra khách gửi trả hàng theo từng sản phẩm.", "Đơn đã giao, số lượng trả hợp lệ.", "Mở chi tiết đơn; chọn sản phẩm; nhập lý do; gửi yêu cầu.", "Yêu cầu trả hàng được tạo ở trạng thái chờ xử lý.", "Cao"],
    ["TC_ADMIN_01", "Đăng nhập admin", "Kiểm tra quyền truy cập khu vực quản trị.", "Tài khoản có vai trò admin.", "Đăng nhập; mở dashboard admin; truy cập API quản trị.", "Admin được phép truy cập; khách thường bị chặn bởi middleware.", "Cao"],
    ["TC_ADMIN_CAT_01", "Quản lý danh mục", "Kiểm tra thêm, sửa, bật/tắt danh mục và ràng buộc sản phẩm.", "Tên danh mục, trạng thái danh mục.", "Tạo danh mục; sửa tên; tắt danh mục; thử xóa danh mục có sản phẩm.", "Danh mục cập nhật đúng; danh mục đang có sản phẩm không bị xóa tùy tiện.", "Cao"],
    ["TC_ADMIN_PRODUCT_01", "Thêm và sửa sản phẩm", "Kiểm tra quản lý sản phẩm, ảnh, SKU và biến thể.", "Tên, danh mục, giá, ảnh, SKU, thuộc tính.", "Tạo sản phẩm; upload ảnh; thêm biến thể; sửa thông tin.", "Sản phẩm, ảnh và biến thể được lưu; SKU/tổ hợp thuộc tính trùng bị từ chối.", "Cao"],
    ["TC_ADMIN_PRODUCT_02", "Xóa hoặc ngừng bán sản phẩm", "Kiểm tra loại bỏ sản phẩm khỏi danh mục bán.", "Mã sản phẩm cần xử lý.", "Admin xóa hoặc chuyển trạng thái sản phẩm.", "Sản phẩm không còn được bán; dữ liệu đơn hàng cũ vẫn giữ được lịch sử.", "Trung bình"],
    ["TC_ADMIN_ORDER_01", "Quản lý đơn hàng", "Kiểm tra danh sách đơn và cập nhật trạng thái.", "Đơn hàng ở các trạng thái pending, confirmed, shipping, delivered, cancelled.", "Mở danh sách; lọc trạng thái; cập nhật trạng thái đơn.", "Danh sách hiển thị đủ thông tin; trạng thái được cập nhật đúng.", "Cao"],
    ["TC_ADMIN_PAYMENT_01", "Xác nhận thanh toán chuyển khoản", "Kiểm tra admin xác nhận hoặc đánh dấu chưa nhận tiền.", "Đơn thanh toán QR.", "Khách báo đã chuyển khoản; admin chọn đã thanh toán hoặc chưa nhận tiền.", "Trạng thái thanh toán và người xác nhận được lưu đúng.", "Cao"],
    ["TC_ADMIN_STOCK_01", "Nhập kho", "Kiểm tra tạo phiếu nhập kho một hoặc nhiều biến thể.", "Mã phiếu, ngày nhập, danh sách biến thể và số lượng.", "Tạo phiếu nhập; thêm nhiều dòng; lưu phiếu.", "Tồn kho tăng đúng và có chi tiết phiếu nhập.", "Cao"],
    ["TC_ADMIN_STOCK_02", "Lịch sử biến động kho", "Kiểm tra mọi biến động tồn kho được ghi nhận.", "Nhập kho, bán hàng, hủy đơn, điều chỉnh, trả hàng.", "Thực hiện các nghiệp vụ làm thay đổi tồn kho; xem lịch sử.", "Mỗi biến động có loại, số lượng trước/sau, người thực hiện và tham chiếu.", "Cao"],
    ["TC_ADMIN_STOCK_03", "Cảnh báo gần hết hàng", "Kiểm tra cảnh báo tồn kho theo ngưỡng cấu hình.", "Biến thể có tồn kho nhỏ hơn hoặc bằng ngưỡng.", "Cập nhật tồn kho thấp; mở màn cảnh báo.", "Sản phẩm sắp hết hàng xuất hiện trong danh sách cảnh báo.", "Trung bình"],
    ["TC_ADMIN_REVIEW_01", "Quản lý đánh giá", "Kiểm tra duyệt và phản hồi đánh giá.", "Đánh giá đang chờ duyệt.", "Admin duyệt đánh giá; nhập phản hồi.", "Đánh giá chuyển trạng thái được duyệt và hiển thị phản hồi admin.", "Trung bình"],
    ["TC_ADMIN_RETURN_01", "Quản lý yêu cầu trả hàng", "Kiểm tra duyệt, nhận hàng trả và cập nhật kho/doanh thu.", "Yêu cầu trả hàng hợp lệ.", "Admin đổi trạng thái nhận hàng; kiểm tra tồn kho và báo cáo.", "Tồn kho tăng lại, doanh thu giảm tương ứng với hàng trả.", "Cao"],
    ["TC_ADMIN_ANN_01", "Quản lý thông báo", "Kiểm tra tạo, upload ảnh, sắp xếp và công bố thông báo.", "Tiêu đề, nội dung, ảnh, trạng thái.", "Tạo thông báo; upload ảnh; đổi thứ tự; công bố.", "Chỉ thông báo đã xuất bản hiển thị công khai; ảnh hợp lệ được lưu.", "Trung bình"],
    ["TC_ADMIN_REPORT_01", "Dashboard doanh thu", "Kiểm tra tổng quan doanh thu và số lượng đơn.", "Tập đơn hàng nhiều trạng thái.", "Mở dashboard; xem thống kê tổng và theo tháng.", "Doanh thu chỉ tính đơn delivered, thống kê trạng thái đúng.", "Cao"],
    ["TC_BIZ_01", "Tính tổng tiền và mã khuyến mãi", "Kiểm tra tạm tính, phí ship, giảm giá và tổng tiền cuối.", "Giỏ hàng, mã SPORT20 hoặc mã hợp lệ.", "Áp dụng mã; đặt hàng; kiểm tra số tiền giảm và tổng tiền.", "Tổng tiền = tạm tính + phí ship - giảm giá; lịch sử dùng mã được ghi nhận.", "Cao"],
    ["TC_BIZ_02", "Tính phí vận chuyển theo khu vực", "Kiểm tra phí ship nội thành, ngoại thành và tỉnh khác.", "Địa chỉ HCM nội thành, HCM ngoại thành, tỉnh khác.", "Gọi tính phí hoặc nhập địa chỉ trên checkout.", "Nội thành miễn phí, ngoại thành 30.000đ, tỉnh khác 50.000đ.", "Cao"],
    ["TC_BIZ_03", "Quy tắc doanh thu", "Kiểm tra đơn chờ xác nhận không tính doanh thu.", "Đơn pending, confirmed, shipping, delivered, cancelled.", "Tạo nhiều đơn; xem báo cáo tổng và báo cáo tháng.", "Chỉ đơn delivered được cộng vào doanh thu.", "Cao"],
    ["TC_BIZ_04", "Cập nhật tồn kho theo nghiệp vụ", "Kiểm tra tồn kho khi bán, hủy, nhập và trả hàng.", "Biến thể có tồn kho, đơn hàng, phiếu nhập, yêu cầu trả.", "Đặt hàng; hủy đơn; nhập kho; nhận hàng trả.", "Tồn kho giảm khi bán, tăng khi hủy/nhập/trả; lịch sử kho được ghi nhận.", "Cao"],
    ["TC_BIZ_05", "Ràng buộc dữ liệu nghiệp vụ", "Kiểm tra chống dữ liệu trùng hoặc không hợp lệ.", "SKU trùng, tổ hợp thuộc tính trùng, mã coupon đã dùng.", "Tạo dữ liệu trùng; dùng lại coupon sau khi hủy đơn.", "Hệ thống từ chối dữ liệu trùng và không cho dùng lại mã đã dùng.", "Cao"],
]


result_rows = [
    ["TC_USER_01", "Đăng ký tài khoản", "Kiểm tra qua mã nguồn AuthController: validate email/số điện thoại duy nhất, tạo KhachHang và GioHang.", "Đạt", "Chưa có test E2E riêng cho form đăng ký; cần kiểm thử thủ công giao diện."],
    ["TC_USER_02", "Đăng nhập và đăng xuất", "Các test chức năng sử dụng xác thực Sanctum thành công; AuthController có xử lý đăng nhập, khóa tài khoản và logout.", "Đạt", "Nên bổ sung test tự động riêng cho sai mật khẩu và logout."],
    ["TC_PRODUCT_01", "Xem, tìm kiếm, lọc sản phẩm", "ProductCatalogTest xác nhận danh mục thể thao, bộ lọc và loại trừ thương hiệu hoạt động.", "Đạt", "Đã kiểm tra API; frontend build thành công."],
    ["TC_PRODUCT_02", "Xem chi tiết sản phẩm", "ProductCatalogTest xác nhận chi tiết sản phẩm trả đủ thông tin biến thể.", "Đạt", "Không phát hiện lỗi trong test tự động."],
    ["TC_CART_01", "Thêm sản phẩm vào giỏ hàng", "PurchaseFlowTest xác nhận thêm biến thể còn hàng vào giỏ thành công.", "Đạt", "Có kiểm tra trường hợp hết hàng bị từ chối."],
    ["TC_CART_02", "Cập nhật và xóa giỏ hàng", "CartController có API cập nhật/xóa; frontend build thành công.", "Đạt", "Chưa có test tự động riêng cho thao tác update/remove UI; cần kiểm thử thủ công bổ sung."],
    ["TC_ORDER_01", "Đặt hàng COD", "PurchaseFlowTest xác nhận tạo đơn COD, giảm tồn kho, xóa giỏ hàng.", "Đạt", "Có kiểm tra địa chỉ hành chính không hợp lệ."],
    ["TC_ORDER_02", "Đặt hàng QR chuyển khoản", "ManualQrPaymentShippingTest xác nhận tạo QR, khách báo đã chuyển khoản và admin xác nhận paid.", "Đạt", "QR URL dùng VietQR."],
    ["TC_ORDER_03", "Lịch sử và chi tiết đơn hàng", "PurchaseFlowTest xác nhận khách xem được đơn của mình và không xem được đơn người khác.", "Đạt", "Đã kiểm tra phân quyền chi tiết đơn."],
    ["TC_REVIEW_01", "Đánh giá sản phẩm", "CommerceFeaturesTest xác nhận chỉ đơn delivered mới đánh giá được, đánh giá chỉ công khai sau khi admin duyệt.", "Đạt", "Đã kiểm tra phản hồi admin."],
    ["TC_RETURN_01", "Gửi yêu cầu trả hàng", "ReturnRequestFeatureTest xác nhận khách tạo yêu cầu trả hợp lệ và từ chối đơn/số lượng không hợp lệ.", "Đạt", "Đã kiểm tra cập nhật kho sau khi admin nhận hàng trả."],
    ["TC_ADMIN_01", "Đăng nhập admin", "Các test admin chạy bằng tài khoản vai_tro=true và truy cập API quản trị thành công.", "Đạt", "Nên bổ sung test riêng cho tài khoản khách truy cập admin bị chặn."],
    ["TC_ADMIN_CAT_01", "Quản lý danh mục", "AdminCatalogManagementTest xác nhận quản lý danh mục với ràng buộc sản phẩm.", "Đạt", "Không phát hiện lỗi trong test tự động."],
    ["TC_ADMIN_PRODUCT_01", "Thêm và sửa sản phẩm", "AdminCatalogManagementTest xác nhận upload ảnh, tạo/cập nhật sản phẩm, biến thể và chống trùng SKU/tổ hợp.", "Đạt", "Không phát hiện lỗi trong test tự động."],
    ["TC_ADMIN_PRODUCT_02", "Xóa hoặc ngừng bán sản phẩm", "Kiểm tra qua mã nguồn/migration: sản phẩm có trạng thái active/inactive/out_of_stock; danh mục có ràng buộc.", "Đạt", "Chưa có test E2E riêng cho xóa sản phẩm đang có đơn."],
    ["TC_ADMIN_ORDER_01", "Quản lý đơn hàng", "AdminOrderListTest và InventoryManagementTest xác nhận danh sách đơn, số lượng sản phẩm và cập nhật trạng thái.", "Đạt", "Hủy đơn có hoàn kho và ghi lịch sử."],
    ["TC_ADMIN_PAYMENT_01", "Xác nhận thanh toán chuyển khoản", "ManualQrPaymentShippingTest xác nhận admin đổi trạng thái thanh toán sang paid, lưu thời điểm và người xác nhận.", "Đạt", "Không phát hiện lỗi trong test tự động."],
    ["TC_ADMIN_STOCK_01", "Nhập kho", "InventoryManagementTest xác nhận nhập một hoặc nhiều biến thể trong một phiếu.", "Đạt", "Số lượng nhập không hợp lệ bị từ chối."],
    ["TC_ADMIN_STOCK_02", "Lịch sử biến động kho", "InventoryManagementTest xác nhận nhập kho, điều chỉnh, bán hàng và hủy đơn đều ghi lịch sử.", "Đạt", "ReturnRequestFeatureTest xác nhận trả hàng cũng ghi biến động kho."],
    ["TC_ADMIN_STOCK_03", "Cảnh báo gần hết hàng", "InventoryManagementTest xác nhận cảnh báo sử dụng ngưỡng cấu hình của biến thể.", "Đạt", "Không phát hiện lỗi trong test tự động."],
    ["TC_ADMIN_REVIEW_01", "Quản lý đánh giá", "CommerceFeaturesTest xác nhận admin duyệt và phản hồi đánh giá.", "Đạt", "Không phát hiện lỗi trong test tự động."],
    ["TC_ADMIN_RETURN_01", "Quản lý yêu cầu trả hàng", "ReturnRequestFeatureTest xác nhận admin nhận hàng trả, tăng tồn kho và cập nhật doanh thu.", "Đạt", "Không phát hiện lỗi trong test tự động."],
    ["TC_ADMIN_ANN_01", "Quản lý thông báo", "AnnouncementImageManagementTest xác nhận upload, tạo, sắp xếp, xóa ảnh; CommerceFeaturesTest xác nhận chỉ published công khai.", "Đạt", "Ảnh không hợp lệ bị từ chối."],
    ["TC_ADMIN_REPORT_01", "Dashboard doanh thu", "AdminRevenueReportTest xác nhận tổng quan và doanh thu tháng chỉ tính đơn delivered.", "Đạt", "Không phát hiện lỗi trong test tự động."],
    ["TC_BIZ_01", "Tính tổng tiền và mã khuyến mãi", "CommerceFeaturesTest xác nhận coupon dùng một lần; OrderController tính tổng theo tạm tính, phí ship và giảm giá.", "Đạt", "Nên bổ sung test riêng cho từng mức giảm giá."],
    ["TC_BIZ_02", "Tính phí vận chuyển theo khu vực", "ManualQrPaymentShippingTest xác nhận nội thành 0đ, ngoại thành 30.000đ, tỉnh khác 50.000đ.", "Đạt", "Đã kiểm tra thiếu địa chỉ chi tiết trả về không hợp lệ."],
    ["TC_BIZ_03", "Quy tắc doanh thu", "AdminRevenueReportTest xác nhận pending/confirmed/shipping/cancelled không cộng doanh thu.", "Đạt", "Tỷ lệ đạt 100% với các test doanh thu hiện có."],
    ["TC_BIZ_04", "Cập nhật tồn kho theo nghiệp vụ", "InventoryManagementTest và ReturnRequestFeatureTest xác nhận tồn kho thay đổi đúng khi bán, hủy, nhập, trả.", "Đạt", "Lịch sử kho được ghi nhận."],
    ["TC_BIZ_05", "Ràng buộc dữ liệu nghiệp vụ", "AdminCatalogManagementTest và CommerceFeaturesTest xác nhận chống trùng SKU/tổ hợp và coupon đã dùng.", "Đạt", "Không phát hiện lỗi trong test tự động."],
]


exception_rows = [
    ["EX_AUTH_01", "Đăng nhập sai mật khẩu", "AuthController kiểm tra Hash và trả lỗi xác thực.", "Không cấp token, hiển thị thông báo email hoặc mật khẩu không đúng.", "Cần kiểm thử thủ công/API bổ sung"],
    ["EX_AUTH_02", "Đăng nhập tài khoản không tồn tại", "Không tìm thấy KhachHang theo email và trả lỗi xác thực.", "Không đăng nhập được, không tiết lộ chi tiết tài khoản tồn tại hay không.", "Cần kiểm thử thủ công/API bổ sung"],
    ["EX_AUTH_03", "Đăng ký trùng email/số điện thoại", "Validate unique trên khach_hang.email và khach_hang.dien_thoai.", "Hệ thống từ chối tạo tài khoản và báo lỗi trường trùng.", "Cần kiểm thử thủ công/API bổ sung"],
    ["EX_AUTH_04", "Bỏ trống trường bắt buộc", "Laravel validation trả lỗi cho tên, email, mật khẩu hoặc dữ liệu thiếu.", "Không lưu dữ liệu thiếu, hiển thị lỗi nhập liệu.", "Cần kiểm thử thủ công/API bổ sung"],
    ["EX_CART_01", "Thêm sản phẩm hết hàng vào giỏ", "CartController kiểm tra so_luong_ton trước khi thêm.", "Trả lỗi 422 và không thêm vào giỏ.", "Đạt qua test tự động"],
    ["EX_CART_02", "Đặt số lượng lớn hơn tồn kho", "CartController kiểm tra số lượng khi thêm và cập nhật giỏ.", "Từ chối thao tác và báo số lượng còn trong kho.", "Đạt qua mã nguồn; cần test UI"],
    ["EX_CART_03", "Xóa sản phẩm khỏi giỏ", "API DELETE hoặc cập nhật số lượng về 0 sẽ xóa dòng giỏ hàng.", "Sản phẩm biến mất khỏi giỏ, tổng tiền tính lại.", "Đạt qua mã nguồn; cần test UI"],
    ["EX_CART_04", "Sản phẩm bị admin xóa/ngừng bán nhưng còn trong giỏ", "OrderController kiểm tra biến thể còn tồn tại và còn hoạt động khi đặt hàng.", "Không cho đặt nếu biến thể không hợp lệ hoặc đã tắt.", "Đạt qua mã nguồn; cần test bổ sung"],
    ["EX_ORDER_01", "Đặt hàng khi chưa đăng nhập", "Các route đặt hàng dùng middleware auth:sanctum.", "Yêu cầu đăng nhập, không tạo đơn hàng.", "Đạt qua mã nguồn"],
    ["EX_ORDER_02", "Đặt hàng thiếu địa chỉ", "OrderController validate address_detail bắt buộc.", "Trả lỗi validate, không tạo đơn.", "Đạt qua test tự động"],
    ["EX_ORDER_03", "Địa chỉ hành chính không hợp lệ", "ShippingPaymentService kiểm tra province/district/ward hợp lệ.", "Trả lỗi 422, không tạo đơn.", "Đạt qua test tự động"],
    ["EX_ORDER_04", "Hủy đơn khi đơn đã giao", "Cần ràng buộc thêm ở nghiệp vụ nếu không cho hủy delivered.", "Không cho hủy đơn đã giao; nếu chưa có ràng buộc thì cần cải thiện.", "Cần kiểm tra thực tế bổ sung"],
    ["EX_ORDER_05", "Xem đơn hàng không thuộc tài khoản hiện tại", "OrderController lọc theo ma_kh của người dùng hiện tại.", "Trả 404 hoặc không hiển thị dữ liệu đơn người khác.", "Đạt qua test tự động"],
    ["EX_PAY_01", "Chọn QR nhưng chưa thanh toán", "Đơn QR tạo với trạng thái pending_payment.", "Đơn không được coi là đã thanh toán cho đến khi khách báo và admin xác nhận.", "Đạt qua test tự động"],
    ["EX_PAY_02", "Khách bấm đã chuyển khoản nhưng admin chưa nhận tiền", "Trạng thái chuyển sang waiting_admin_confirmation.", "Đơn chờ admin kiểm tra, chưa chuyển sang paid.", "Đạt qua test tự động"],
    ["EX_PAY_03", "Admin xác nhận nhầm trạng thái thanh toán", "Admin có thể đánh dấu paid hoặc payment_not_received.", "Trạng thái thanh toán được cập nhật theo lựa chọn; cần quy trình đối soát tránh thao tác nhầm.", "Đạt qua mã nguồn; cần kiểm thử thủ công"],
    ["EX_PAY_04", "Đơn chờ thanh toán không tính doanh thu", "Báo cáo doanh thu dựa trên trạng thái đơn delivered.", "Đơn chưa hoàn tất không cộng vào doanh thu.", "Đạt qua test tự động"],
    ["EX_SHIP_01", "Chưa chọn tỉnh/thành phố", "ShippingPaymentService trả valid=false nếu province_type không hợp lệ.", "Không tính phí, yêu cầu chọn tỉnh/thành.", "Đạt qua mã nguồn"],
    ["EX_SHIP_02", "Chưa chọn quận/huyện", "Với HCM/Hà Nội, service yêu cầu district_code.", "Không tính phí, yêu cầu chọn quận/huyện.", "Đạt qua mã nguồn"],
    ["EX_SHIP_03", "Chưa chọn phường/xã", "Với HCM/Hà Nội, service yêu cầu ward_code.", "Không tính phí, yêu cầu chọn phường/xã.", "Đạt qua mã nguồn"],
    ["EX_SHIP_04", "Chọn tỉnh khác", "Service cho phép province_type=other và bỏ qua quận/huyện.", "Tính phí 50.000đ, khóa lựa chọn quận/huyện trên giao diện.", "Đạt qua test tự động; cần kiểm tra UI khóa trường"],
    ["EX_SHIP_05", "Chọn nội thành", "Shipping zone inner_city trả phí 0đ.", "Miễn phí vận chuyển.", "Đạt qua test tự động"],
    ["EX_SHIP_06", "Chọn ngoại thành", "Shipping zone suburban trả phí 30.000đ.", "Tính phí vận chuyển 30.000đ.", "Đạt qua test tự động"],
    ["EX_STOCK_01", "Nhập kho số lượng âm hoặc bằng 0", "Validation yêu cầu quantity tối thiểu 1.", "Từ chối phiếu nhập, không thay đổi tồn kho.", "Đạt qua test tự động"],
    ["EX_STOCK_02", "Bán vượt quá tồn kho", "CartController và OrderController kiểm tra tồn kho trước khi thêm/đặt.", "Không cho thêm hoặc đặt hàng vượt tồn.", "Đạt qua test tự động"],
    ["EX_STOCK_03", "Cập nhật tồn kho nhưng không ghi lịch sử", "InventoryService luôn tạo LichSuBienDongKho khi thay đổi tồn.", "Mọi thay đổi tồn có lịch sử trước/sau và tham chiếu.", "Đạt qua test tự động"],
    ["EX_STOCK_04", "Sản phẩm dưới ngưỡng nhưng không cảnh báo", "API cảnh báo lọc theo so_luong_ton <= nguong_canh_bao_ton.", "Biến thể dưới ngưỡng xuất hiện trong cảnh báo.", "Đạt qua test tự động"],
    ["EX_REVIEW_01", "Đánh giá khi đơn chưa hoàn thành", "ReviewController từ chối nếu đơn chưa delivered.", "Không tạo đánh giá.", "Đạt qua test tự động"],
    ["EX_REVIEW_02", "Đánh giá trùng sản phẩm trong cùng đơn", "Migration có unique ma_dh + ma_sp.", "Không cho tạo đánh giá trùng cho cùng đơn và sản phẩm.", "Đạt qua mã nguồn; cần test bổ sung"],
    ["EX_RETURN_01", "Trả hàng với đơn chưa giao thành công", "ReturnRequestController chỉ cho đơn delivered.", "Từ chối yêu cầu trả hàng.", "Đạt qua test tự động"],
    ["EX_RETURN_02", "Trả số lượng lớn hơn đã mua", "Controller tính số lượng còn có thể trả và validate.", "Từ chối số lượng vượt quá đã mua/còn lại.", "Đạt qua test tự động"],
]


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("CHƯƠNG 4. THỬ NGHIỆM")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16)

    add_body(doc, "Chương này trình bày kế hoạch và kết quả thử nghiệm hệ thống bán hàng TienProSport. Nội dung thử nghiệm tập trung vào các nhóm chức năng chính của khách hàng, quản trị viên, nghiệp vụ tính toán, thanh toán, vận chuyển, kho hàng, đánh giá và trả hàng.")
    add_body(doc, "backend API Laravel, các test chức năng hiện có, kiểm tra build frontend React/Vite và đối chiếu mã nguồn với các luồng nghiệp vụ chính.", "Phạm vi kiểm thử: ")
    add_body(doc, "PHP 8.3.14, PHPUnit 12.5.30, Laravel feature tests, Vite 6.3.5. Bộ kiểm thử backend chạy ngày 26/06/2026 đạt 32/32 test với 236 assertion. Frontend build production thành công; Vite chỉ cảnh báo kích thước bundle lớn hơn 500 kB, không làm hỏng build.", "Môi trường kiểm thử: ")
    add_body(doc, "Các dòng ghi 'Đạt qua test tự động' dựa trên kết quả chạy php artisan test. Các dòng ghi 'Đạt qua mã nguồn' được xác nhận bằng controller/service hiện có nhưng vẫn nên kiểm thử giao diện thủ công trước khi nghiệm thu cuối.", "Cách ghi nhận kết quả: ")

    add_heading(doc, "4.1. CÁC KỊCH BẢN THỬ NGHIỆM", 1)
    add_body(doc, "Bảng sau tổng hợp các kịch bản thử nghiệm chính cho hệ thống TienProSport. Các kịch bản được nhóm theo chức năng khách hàng, quản trị viên và nghiệp vụ lõi của hệ thống.")
    add_table(
        doc,
        ["Mã kịch bản", "Tên kịch bản", "Mục tiêu kiểm thử", "Dữ liệu đầu vào", "Các bước thực hiện", "Kết quả mong đợi", "Ưu tiên"],
        scenario_rows,
        [0.72, 1.25, 1.55, 1.50, 2.35, 2.05, 0.72],
        7.0,
    )

    add_heading(doc, "4.2. KẾT QUẢ THỬ NGHIỆM CÁC KỊCH BẢN", 1)
    add_body(doc, "Kết quả thử nghiệm được tổng hợp từ bộ test tự động của backend, kiểm tra build frontend và đối chiếu mã nguồn ở các chức năng chưa có test giao diện. Không phát hiện lỗi nghiêm trọng trong các test tự động đã chạy.")
    add_table(
        doc,
        ["Mã kịch bản", "Tên kịch bản", "Kết quả thực tế", "Trạng thái", "Ghi chú"],
        result_rows,
        [0.85, 1.55, 4.25, 0.85, 3.00],
        7.8,
    )
    add_body(doc, "Bộ kiểm thử tự động hiện có đạt 32/32 test, tương đương 100% số test được chạy, với 236 assertion đạt. Các chức năng ổn định gồm danh mục, sản phẩm, giỏ hàng, đặt hàng COD, thanh toán QR, phí vận chuyển, doanh thu, tồn kho, đánh giá, thông báo và trả hàng. Frontend build production thành công, cho thấy mã giao diện có thể biên dịch. Các phần nên cải thiện là bổ sung test E2E thao tác giao diện cho đăng ký/đăng nhập/đăng xuất, cập nhật-xóa giỏ hàng, xóa/ngừng bán sản phẩm và quy tắc không cho hủy đơn đã giao nếu đây là yêu cầu nghiệp vụ bắt buộc.", "Nhận xét tổng quan: ")
    add_body(doc, "Hệ thống đáp ứng tốt các yêu cầu nghiệp vụ chính của website bán hàng. Các luồng có rủi ro tài chính như thanh toán QR, tính doanh thu, trừ/hoàn tồn kho và trả hàng đã có test tự động xác nhận. Trước khi nghiệm thu chính thức, nên chạy thêm kiểm thử thủ công trên trình duyệt với dữ liệu thật hoặc môi trường staging.", "Đánh giá mức độ đáp ứng: ")

    add_heading(doc, "4.3. XỬ LÝ CÁC KỊCH BẢN NGOẠI LỆ", 1)
    add_body(doc, "Phần này trình bày các tình huống nhập liệu sai, thao tác sai hoặc trạng thái bất thường. Trạng thái kiểm thử phản ánh mức đã kiểm chứng trong lần kiểm tra hiện tại.")
    add_table(
        doc,
        ["Mã ngoại lệ", "Tình huống ngoại lệ", "Cách xử lý của hệ thống", "Kết quả mong đợi", "Trạng thái kiểm thử"],
        exception_rows,
        [0.85, 2.20, 3.05, 2.75, 1.65],
        7.8,
    )
    add_body(doc, "Qua kiểm thử tự động và đối chiếu mã nguồn, hệ thống TienProSport đã đáp ứng phần lớn chức năng cốt lõi của website bán hàng. Các nghiệp vụ quan trọng như đặt hàng, thanh toán QR, tính phí vận chuyển, quản lý kho, doanh thu, đánh giá và trả hàng được kiểm tra rõ ràng. Một số kịch bản giao diện cần tiếp tục kiểm thử thủ công để hoàn thiện độ bao phủ trước khi triển khai thực tế.", "Kết luận chương: ")

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("TienProSport - Báo cáo thử nghiệm")
        footer_run.font.name = "Times New Roman"
        footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        footer_run.font.size = Pt(10)

    doc.save(OUT)


if __name__ == "__main__":
    build()
