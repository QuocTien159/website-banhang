from __future__ import annotations

import math
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


INPUT = Path(r"C:\Users\ACER\Downloads\testLVTN_Chuong5.docx")
OUTPUT = Path(r"C:\Users\ACER\Downloads\testLVTN_Chuong5_UseCase_v2.docx")
OUT_DIR = Path(r"C:\Users\ACER\Downloads\banhangcodex\docs\usecase_diagrams")


DIAGRAMS = [
    {
        "heading": "Use case Ä‘Äƒng kÃ½",
        "file": "uc_dang_ky.png",
        "actor": "NgÆ°á»i dÃ¹ng",
        "main": "ÄÄƒng kÃ½",
        "include": ["XÃ¡c thá»±c dá»¯ liá»‡u nháº­p"],
        "extend": ["ÄÄƒng kÃ½ báº±ng tÃ i khoáº£n Google"],
    },
    {
        "heading": "Use case Ä‘Äƒng nháº­p",
        "file": "uc_dang_nhap.png",
        "actor": "NgÆ°á»i dÃ¹ng",
        "main": "ÄÄƒng nháº­p",
        "include": ["XÃ¡c thá»±c thÃ´ng tin", "Kiá»ƒm tra vai trÃ²"],
        "extend": ["ÄÄƒng nháº­p báº±ng Google"],
    },
    {
        "heading": "Use case chá»‰nh sá»­a thÃ´ng tin cÃ¡ nhÃ¢n",
        "file": "uc_chinh_sua_thong_tin.png",
        "actor": "NgÆ°á»i dÃ¹ng",
        "main": "Chá»‰nh sá»­a thÃ´ng tin cÃ¡ nhÃ¢n",
        "include": ["XÃ¡c thá»±c dá»¯ liá»‡u nháº­p", "Cáº­p nháº­t thÃ´ng tin"],
        "extend": ["Äá»•i máº­t kháº©u"],
    },
    {
        "heading": "Use case tÃ¬m kiáº¿m sáº£n pháº©m",
        "file": "uc_tim_kiem_san_pham.png",
        "actor": "KhÃ¡ch hÃ ng",
        "main": "TÃ¬m kiáº¿m sáº£n pháº©m",
        "include": ["Nháº­p tá»« khÃ³a tÃ¬m kiáº¿m", "Hiá»ƒn thá»‹ káº¿t quáº£"],
        "extend": ["Lá»c theo danh má»¥c / thuá»™c tÃ­nh"],
    },
    {
        "heading": "Use case quáº£n lÃ½ giá» hÃ ng",
        "file": "uc_quan_ly_gio_hang.png",
        "actor": "KhÃ¡ch hÃ ng",
        "main": "Quáº£n lÃ½ giá» hÃ ng",
        "include": ["ThÃªm sáº£n pháº©m vÃ o giá»", "Cáº­p nháº­t sá»‘ lÆ°á»£ng"],
        "extend": ["XÃ³a sáº£n pháº©m khá»i giá»"],
    },
    {
        "heading": "Use case thanh toÃ¡n Ä‘Æ¡n hÃ ng",
        "file": "uc_thanh_toan_don_hang.png",
        "actor": "KhÃ¡ch hÃ ng",
        "main": "Thanh toÃ¡n Ä‘Æ¡n hÃ ng",
        "include": ["XÃ¡c nháº­n giá» hÃ ng", "TÃ­nh phÃ­ váº­n chuyá»ƒn"],
        "extend": ["Thanh toÃ¡n COD", "Thanh toÃ¡n QR chuyá»ƒn khoáº£n"],
    },
    {
        "heading": "Use case yÃªu cáº§u tráº£ hÃ ng / hoÃ n tiá»n",
        "file": "uc_yeu_cau_tra_hang.png",
        "actor": "KhÃ¡ch hÃ ng",
        "main": "YÃªu cáº§u tráº£ hÃ ng / hoÃ n tiá»n",
        "include": ["Kiá»ƒm tra Ä‘Æ¡n Ä‘Ã£ giao", "Nháº­p lÃ½ do vÃ  sá»‘ lÆ°á»£ng"],
        "extend": ["Táº£i áº£nh minh chá»©ng"],
    },
    {
        "heading": "Use case xá»­ lÃ½ yÃªu cáº§u tráº£ hÃ ng",
        "file": "uc_xu_ly_tra_hang.png",
        "actor": "Quáº£n trá»‹ viÃªn",
        "main": "Xá»­ lÃ½ yÃªu cáº§u tráº£ hÃ ng",
        "include": ["Xem chi tiáº¿t yÃªu cáº§u", "Cáº­p nháº­t tráº¡ng thÃ¡i"],
        "extend": ["Nháº­p láº¡i kho / hoÃ n tiá»n"],
    },
    {
        "heading": "Use case quáº£n lÃ½ danh má»¥c",
        "file": "uc_quan_ly_danh_muc.png",
        "actor": "Quáº£n trá»‹ viÃªn",
        "main": "Quáº£n lÃ½ danh má»¥c",
        "include": ["ThÃªm / sá»­a danh má»¥c", "Kiá»ƒm tra rÃ ng buá»™c sáº£n pháº©m"],
        "extend": ["KhÃ³a hoáº·c áº©n danh má»¥c"],
    },
    {
        "heading": "Use case quáº£n lÃ½ sáº£n pháº©m",
        "file": "uc_quan_ly_san_pham.png",
        "actor": "Quáº£n trá»‹ viÃªn",
        "main": "Quáº£n lÃ½ sáº£n pháº©m",
        "include": ["Quáº£n lÃ½ biáº¿n thá»ƒ", "Quáº£n lÃ½ hÃ¬nh áº£nh"],
        "extend": ["Cáº£nh bÃ¡o trÃ¹ng SKU"],
    },
    {
        "heading": "Use case quáº£n lÃ½ Ä‘Æ¡n hÃ ng",
        "file": "uc_quan_ly_don_hang.png",
        "actor": "Quáº£n trá»‹ viÃªn",
        "main": "Quáº£n lÃ½ Ä‘Æ¡n hÃ ng",
        "include": ["Xem danh sÃ¡ch Ä‘Æ¡n", "Cáº­p nháº­t tráº¡ng thÃ¡i Ä‘Æ¡n"],
        "extend": ["XÃ¡c nháº­n thanh toÃ¡n QR"],
    },
    {
        "heading": "Use case quáº£n lÃ½ khuyáº¿n mÃ£i",
        "file": "uc_quan_ly_khuyen_mai.png",
        "actor": "Quáº£n trá»‹ viÃªn",
        "main": "Quáº£n lÃ½ khuyáº¿n mÃ£i",
        "include": ["Táº¡o / sá»­a mÃ£ giáº£m giÃ¡", "Kiá»ƒm tra Ä‘iá»u kiá»‡n Ã¡p dá»¥ng"],
        "extend": ["Ngá»«ng kÃ­ch hoáº¡t mÃ£"],
    },
    {
        "heading": "Use case quáº£n lÃ½ tá»“n kho",
        "file": "uc_quan_ly_ton_kho.png",
        "actor": "Quáº£n trá»‹ viÃªn",
        "main": "Quáº£n lÃ½ tá»“n kho",
        "include": ["Nháº­p kho", "Ghi lá»‹ch sá»­ biáº¿n Ä‘á»™ng"],
        "extend": ["Cáº£nh bÃ¡o gáº§n háº¿t hÃ ng"],
    },
    {
        "heading": "Use case bÃ¡o cÃ¡o thá»‘ng kÃª",
        "file": "uc_bao_cao_thong_ke.png",
        "actor": "Quáº£n trá»‹ viÃªn",
        "main": "BÃ¡o cÃ¡o thá»‘ng kÃª",
        "include": ["Thá»‘ng kÃª doanh thu", "Thá»‘ng kÃª Ä‘Æ¡n hÃ ng"],
        "extend": ["Lá»c theo thá»i gian"],
    },
]


def load_font(size: int, bold: bool = False):
    candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf"),
        Path(r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


FONT = load_font(24)
BOLD_FONT = load_font(25, True)
SMALL_FONT = load_font(18)


def text_size(draw: ImageDraw.ImageDraw, text: str, font) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def wrap_text_to_width(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if text_size(draw, candidate, font)[0] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def centered_text(draw: ImageDraw.ImageDraw, center: tuple[int, int], text: str, font, max_width: int, fill="black"):
    lines = wrap_text_to_width(draw, text, font, max_width)
    heights = [text_size(draw, line, font)[1] for line in lines]
    total_h = sum(heights) + max(0, len(lines) - 1) * 4
    y = center[1] - total_h / 2
    for line, h in zip(lines, heights):
        w, _ = text_size(draw, line, font)
        draw.text((center[0] - w / 2, y), line, font=font, fill=fill)
        y += h + 4


def ellipse(draw, box, text, font=BOLD_FONT, shadow=True):
    if shadow:
        shadow_box = (box[0] + 16, box[1] + 15, box[2] + 16, box[3] + 15)
        draw.ellipse(shadow_box, fill=(232, 232, 232), outline=None)
    draw.ellipse(box, fill="white", outline="black", width=3)
    cx = (box[0] + box[2]) // 2
    cy = (box[1] + box[3]) // 2
    centered_text(draw, (cx, cy), text, font, max_width=(box[2] - box[0] - 30))


def line(draw, start, end, fill="black", width=3, dashed=False, arrow=True):
    if dashed:
        draw_dashed_line(draw, start, end, fill=fill, width=width, dash=13, gap=9)
    else:
        draw.line([start, end], fill=fill, width=width)
    if arrow:
        draw_arrow_head(draw, start, end, fill=fill, width=width)


def draw_dashed_line(draw, start, end, fill="black", width=2, dash=10, gap=6):
    x1, y1 = start
    x2, y2 = end
    length = math.hypot(x2 - x1, y2 - y1)
    if length == 0:
        return
    vx = (x2 - x1) / length
    vy = (y2 - y1) / length
    pos = 0
    while pos < length:
        segment_end = min(pos + dash, length)
        sx = x1 + vx * pos
        sy = y1 + vy * pos
        ex = x1 + vx * segment_end
        ey = y1 + vy * segment_end
        draw.line([(sx, sy), (ex, ey)], fill=fill, width=width)
        pos += dash + gap


def draw_arrow_head(draw, start, end, fill="black", width=3):
    x1, y1 = start
    x2, y2 = end
    angle = math.atan2(y2 - y1, x2 - x1)
    size = 18
    left = (
        x2 - size * math.cos(angle - math.pi / 7),
        y2 - size * math.sin(angle - math.pi / 7),
    )
    right = (
        x2 - size * math.cos(angle + math.pi / 7),
        y2 - size * math.sin(angle + math.pi / 7),
    )
    draw.line([left, (x2, y2), right], fill=fill, width=width)


def actor(draw, x: int, y: int, label: str):
    # x, y are the top of the head.
    draw.ellipse((x - 18, y, x + 18, y + 36), outline="black", width=3)
    draw.line((x, y + 36, x, y + 103), fill="black", width=3)
    draw.line((x - 45, y + 62, x + 45, y + 62), fill="black", width=3)
    draw.line((x, y + 103, x - 42, y + 158), fill="black", width=3)
    draw.line((x, y + 103, x + 42, y + 158), fill="black", width=3)
    centered_text(draw, (x, y + 190), label, BOLD_FONT, 170)


def edge_point_ellipse(box, from_point):
    cx = (box[0] + box[2]) / 2
    cy = (box[1] + box[3]) / 2
    rx = (box[2] - box[0]) / 2
    ry = (box[3] - box[1]) / 2
    dx = from_point[0] - cx
    dy = from_point[1] - cy
    if dx == 0 and dy == 0:
        return int(cx), int(cy)
    scale = 1 / math.sqrt((dx * dx) / (rx * rx) + (dy * dy) / (ry * ry))
    return int(cx + dx * scale), int(cy + dy * scale)


def relation_label(draw, start, end, text):
    mx = (start[0] + end[0]) / 2
    my = (start[1] + end[1]) / 2
    w, h = text_size(draw, text, SMALL_FONT)
    draw.rectangle((mx - w / 2 - 6, my - h / 2 - 4, mx + w / 2 + 6, my + h / 2 + 4), fill="white")
    draw.text((mx - w / 2, my - h / 2), text, font=SMALL_FONT, fill="black")


def draw_diagram(spec: dict, out_path: Path):
    width, height = 1280, 430
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    main_box = (250, 173, 500, 253)
    include_boxes = []
    extend_boxes = []

    include_count = len(spec["include"])
    include_start_y = 35 if include_count == 1 else 25
    for i, text in enumerate(spec["include"]):
        include_boxes.append((650, include_start_y + i * 92, 1110, include_start_y + 70 + i * 92))

    extend_count = len(spec["extend"])
    extend_start_y = 285 if extend_count == 1 else 245
    for i, text in enumerate(spec["extend"]):
        extend_boxes.append((650, extend_start_y + i * 82, 1190, extend_start_y + 70 + i * 82))

    actor(draw, 95, 110, spec["actor"])
    ellipse(draw, main_box, spec["main"])

    # Actor association.
    line(draw, (180, 213), edge_point_ellipse(main_box, (180, 213)), width=3, dashed=False)

    for box, text in zip(include_boxes, spec["include"]):
        ellipse(draw, box, text)
        start = edge_point_ellipse(main_box, ((box[0] + box[2]) / 2, (box[1] + box[3]) / 2))
        end = edge_point_ellipse(box, ((main_box[0] + main_box[2]) / 2, (main_box[1] + main_box[3]) / 2))
        line(draw, start, end, width=2, dashed=True)
        relation_label(draw, start, end, "Â«includeÂ»")

    for box, text in zip(extend_boxes, spec["extend"]):
        ellipse(draw, box, text)
        start = edge_point_ellipse(box, ((main_box[0] + main_box[2]) / 2, (main_box[1] + main_box[3]) / 2))
        end = edge_point_ellipse(main_box, ((box[0] + box[2]) / 2, (box[1] + box[3]) / 2))
        line(draw, start, end, width=2, dashed=True)
        relation_label(draw, start, end, "Â«extendÂ»")

    image.save(out_path)


def insert_paragraph_after(anchor, paragraph):
    anchor._p.addnext(paragraph._p)


def set_caption_style(paragraph, text: str):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True


def add_picture_after(doc: Document, heading_para, image_path: Path, caption: str):
    image_para = doc.add_paragraph()
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.paragraph_format.space_after = Pt(2)
    image_para.add_run().add_picture(str(image_path), width=Inches(6.3))

    caption_para = doc.add_paragraph()
    set_caption_style(caption_para, caption)

    insert_paragraph_after(heading_para, image_para)
    insert_paragraph_after(image_para, caption_para)


def normalize(text: str) -> str:
    return " ".join(text.split()).strip().lower()


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    image_paths = {}
    for spec in DIAGRAMS:
        path = OUT_DIR / spec["file"]
        draw_diagram(spec, path)
        image_paths[spec["heading"]] = path

    doc = Document(INPUT)
    heading_map = {}
    for para in doc.paragraphs:
        text = " ".join(para.text.split()).strip()
        if para.style.name == "Heading 4" and normalize(text) in {normalize(item["heading"]) for item in DIAGRAMS}:
            heading_map[normalize(text)] = para

    missing = []
    for spec in DIAGRAMS:
        key = normalize(spec["heading"])
        if key not in heading_map:
            missing.append(spec["heading"])
    if missing:
        raise RuntimeError("KhÃ´ng tÃ¬m tháº¥y heading: " + ", ".join(missing))

    # Insert from bottom to top so XML moves do not disturb following anchors.
    for idx, spec in reversed(list(enumerate(DIAGRAMS, start=1))):
        heading_para = heading_map[normalize(spec["heading"])]
        caption = f"HÃ¬nh 3.{idx}. SÆ¡ Ä‘á»“ {spec['heading'].lower()}"
        add_picture_after(doc, heading_para, image_paths[spec["heading"]], caption)

    doc.save(OUTPUT)
    print(OUTPUT)
    print(f"Inserted {len(DIAGRAMS)} diagrams")


if __name__ == "__main__":
    main()

