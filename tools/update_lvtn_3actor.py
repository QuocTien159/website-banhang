from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


INPUT = Path(r"C:\Users\ACER\Downloads\LVTN.docx")
OUTPUT = Path(r"C:\Users\ACER\Downloads\LVTN_CapNhat_3Actor.docx")

FONT_NAME = "Times New Roman"
BODY_SIZE = 13


TARGETED_PARAGRAPH_TEXT = {
    28: "Vì vậy, việc xây dựng một hệ thống website bán hàng cho TienProSport không chỉ nhằm mục đích giới thiệu sản phẩm mà còn hỗ trợ quản lý quy trình bán hàng, phân quyền người dùng, xử lý đơn hàng, quản lý tồn kho và theo dõi doanh thu một cách tập trung.",
    29: "Từ nhu cầu đó, đề tài hướng đến việc xây dựng website bán hàng TienProSport với các chức năng mua hàng trực tuyến, quản lý sản phẩm, đơn hàng, khách hàng, nhân viên, tồn kho và báo cáo doanh thu. Hệ thống hiện tại tính phí vận chuyển theo khu vực, hỗ trợ thanh toán khi nhận hàng và thanh toán QR chuyển khoản thủ công thay vì tích hợp trực tiếp với đơn vị vận chuyển bên thứ ba.",
    31: "Những khó khăn này ảnh hưởng đến hiệu quả vận hành và trải nghiệm khách hàng. Vì vậy, hệ thống cần hỗ trợ quản lý đơn hàng, tồn kho, thanh toán và phí vận chuyển nội bộ rõ ràng, đồng thời đặt nền tảng để có thể tích hợp các dịch vụ giao vận trong các phiên bản sau.",
    32: "Đề tài được thực hiện nhằm xây dựng website bán hàng TienProSport với mô hình gồm khách hàng, nhân viên và quản trị viên. Trong đó, khách hàng thực hiện các thao tác mua hàng; nhân viên xử lý nghiệp vụ vận hành như sản phẩm, đơn hàng, tồn kho, đánh giá và trả hàng; quản trị viên quản lý hệ thống, nhân viên, cấu hình và báo cáo.",
    42: "Mục tiêu của đề tài là xây dựng website bán hàng TienProSport giúp khách hàng tra cứu, đặt mua sản phẩm thể thao trực tuyến; đồng thời hỗ trợ nhân viên và quản trị viên quản lý sản phẩm, đơn hàng, tồn kho, khách hàng, nhân viên, thanh toán, vận chuyển nội bộ và báo cáo doanh thu trên một hệ thống tập trung.",
    49: "Xây dựng hệ thống bán hàng trực tuyến cho phép khách hàng xem sản phẩm, tìm kiếm, lọc sản phẩm, xem chi tiết, thêm vào giỏ hàng, đặt hàng, thanh toán COD hoặc QR chuyển khoản thủ công và theo dõi lịch sử đơn hàng.",
    50: "Xây dựng phân hệ vận hành cho nhân viên và quản trị viên nhằm quản lý sản phẩm, danh mục, thuộc tính/biến thể, đơn hàng, tồn kho, nhập kho, cảnh báo gần hết hàng, đánh giá sản phẩm và yêu cầu trả hàng.",
    51: "Xây dựng cơ chế phân quyền theo vai trò gồm khách hàng, nhân viên và quản trị viên. Khách hàng chỉ sử dụng chức năng mua hàng; nhân viên thực hiện nghiệp vụ vận hành hằng ngày; quản trị viên có quyền quản lý toàn bộ hệ thống, nhân viên, khách hàng, cấu hình và báo cáo.",
    52: "Hệ thống hiện tại tính phí vận chuyển theo khu vực và hỗ trợ tạo mã QR chuyển khoản để khách hàng thanh toán thủ công. Các chức năng tích hợp API vận chuyển, đối soát ngân hàng tự động và webhook trạng thái giao hàng được xem là hướng phát triển trong tương lai.",
    57: "Luận văn tập trung vào việc xây dựng website bán hàng TienProSport và các nghiệp vụ cốt lõi của hệ thống gồm mua hàng, quản lý sản phẩm, quản lý đơn hàng, quản lý kho, phân quyền người dùng, thanh toán COD/QR thủ công, tính phí vận chuyển theo khu vực và thống kê doanh thu.",
    59: "Phạm vi nghiên cứu bao gồm các nghiệp vụ đặt hàng, thanh toán, quản lý kho, quản lý sản phẩm, xử lý đơn hàng, phân quyền theo vai trò và báo cáo thống kê. Các tích hợp tự động với đơn vị vận chuyển hoặc cổng thanh toán bên thứ ba chưa thuộc phạm vi triển khai hiện tại.",
    62: "Xây dựng chức năng cấu hình phí vận chuyển theo khu vực và thông tin thanh toán QR/ngân hàng trong hệ thống quản trị.",
    63: "Xây dựng luồng xử lý đơn hàng nội bộ từ đặt hàng, xác nhận, chuẩn bị hàng, giao hàng đến hoàn tất hoặc hủy/trả hàng; việc đồng bộ tự động trạng thái từ đơn vị vận chuyển bên thứ ba được định hướng cho giai đoạn phát triển tiếp theo.",
    67: "Website TienProSport được xây dựng theo mô hình bán hàng trực tuyến nội bộ, trong đó hệ thống tập trung quản lý sản phẩm, đơn hàng, kho, khách hàng, nhân viên và báo cáo. Các dịch vụ giao vận bên ngoài như GHN, GHTK hoặc Viettel Post có thể được tích hợp thêm trong tương lai.",
    68: "Giới hạn tính năng đối với Khách hàng: Khách hàng có thể đăng ký, đăng nhập, xem danh sách sản phẩm, tìm kiếm/lọc sản phẩm, xem chi tiết sản phẩm, thêm sản phẩm vào giỏ hàng, đặt hàng, lựa chọn phương thức thanh toán COD hoặc QR chuyển khoản thủ công, xem lịch sử đơn hàng, đánh giá sản phẩm sau khi nhận hàng và gửi yêu cầu trả hàng nếu đơn hàng đủ điều kiện.",
    69: "Giới hạn tính năng đối với Quản trị viên: Quản trị viên có quyền cao nhất trong hệ thống, bao gồm quản lý nhân viên, khách hàng, sản phẩm, danh mục, thuộc tính/biến thể, đơn hàng, tồn kho, nhập kho, đánh giá, trả hàng, chương trình khuyến mãi, thông báo, cấu hình thanh toán - vận chuyển và xem dashboard/báo cáo thống kê.",
    70: "Đề tài không đi sâu vào tích hợp API giao vận hoặc thanh toán tự động. Hệ thống hiện tại xử lý phí vận chuyển theo khu vực, thanh toán QR chuyển khoản thủ công và ghi nhận doanh thu theo các đơn hàng hợp lệ.",
    108: "Từ mô hình Coolmate có thể rút ra bài học rằng hệ thống thương mại điện tử cần tách rõ vai trò người dùng, tối ưu trải nghiệm mua hàng, quản lý vận hành tập trung và có khả năng mở rộng sang tích hợp giao vận hoặc thanh toán tự động khi cần.",
    115: "Laravel API cung cấp cơ chế định tuyến, middleware và phân quyền theo vai trò. Trong hệ thống TienProSport, backend kiểm tra quyền truy cập API theo các vai trò customer, staff và admin, giúp khách hàng không truy cập được trang quản trị, nhân viên chỉ sử dụng các chức năng vận hành được phân quyền, còn quản trị viên có quyền quản lý hệ thống.",
    120: "Giao diện quản trị/vận hành cho Nhân viên và Quản trị viên: Giao diện này hỗ trợ quản lý sản phẩm, danh mục, thuộc tính, đơn hàng, tồn kho, nhập kho, đánh giá, trả hàng và thống kê. Các menu cấp cao như quản lý nhân viên, khách hàng, cấu hình thanh toán - vận chuyển, thông báo và báo cáo đầy đủ chỉ dành cho quản trị viên.",
    127: "Hệ thống Website có các quy trình nghiệp vụ được phân chia rõ ràng theo 3 đối tượng: Khách hàng, Nhân viên và Quản trị viên.",
    205: "Từ các quy trình trên, có thể thấy hệ thống website TienProSport được chia thành ba nhóm chức năng chính: nhóm chức năng dành cho Khách hàng, nhóm chức năng vận hành dành cho Nhân viên và nhóm chức năng quản trị hệ thống dành cho Quản trị viên.",
    386: "Hệ thống TienProSport áp dụng phân quyền theo vai trò gồm khách hàng, nhân viên và quản trị viên. Frontend chỉ hiển thị menu phù hợp với vai trò đăng nhập, đồng thời backend sử dụng middleware để kiểm tra quyền khi gọi API.",
    392: "Nhân viên được phép truy cập các chức năng vận hành như sản phẩm, danh mục, đơn hàng, kho, trả hàng và đánh giá ở mức được phân quyền. Nhân viên không được truy cập các chức năng quản lý nhân viên, phân quyền, cấu hình hệ thống quan trọng hoặc báo cáo đầy đủ dành riêng cho quản trị viên.",
    397: "Quản trị viên có quyền cao nhất trong hệ thống, có thể quản lý nhân viên, khách hàng, cấu hình thanh toán - vận chuyển, thông báo, chương trình khuyến mãi, báo cáo và thực hiện các chức năng vận hành khi cần.",
    398: "Cơ chế này giúp hệ thống hạn chế truy cập trái phép, tách biệt rõ trách nhiệm giữa khách hàng, nhân viên và quản trị viên, đồng thời phù hợp với cách tổ chức mã nguồn hiện tại của TienProSport.",
    404: "Kết quả thử nghiệm cho thấy hệ thống đáp ứng các nghiệp vụ chính của website bán hàng gồm mua hàng, quản lý sản phẩm, xử lý đơn hàng, quản lý tồn kho, thanh toán COD/QR thủ công, tính phí vận chuyển theo khu vực và phân quyền theo ba vai trò.",
    410: "Hệ thống đã hỗ trợ ba nhóm người dùng gồm khách hàng, nhân viên và quản trị viên. Khách hàng thực hiện quy trình mua hàng; nhân viên đảm nhiệm nghiệp vụ vận hành hằng ngày; quản trị viên quản lý toàn bộ hệ thống, nhân viên, cấu hình và báo cáo.",
    411: "Các chức năng quản lý sản phẩm, danh mục, thuộc tính, đơn hàng, kho, nhập kho, đánh giá và trả hàng được phân quyền cho nhân viên và quản trị viên. Các chức năng quản lý nhân viên, khách hàng, cấu hình thanh toán - vận chuyển, thông báo, khuyến mãi và báo cáo đầy đủ được giới hạn cho quản trị viên.",
    412: "Doanh thu trong dashboard và báo cáo được ghi nhận dựa trên các đơn hàng hợp lệ, không tính các đơn đang chờ xác nhận hoặc chờ thanh toán. Thanh toán QR hiện được xử lý theo hình thức chuyển khoản thủ công và cần nhân viên/quản trị viên xác nhận trạng thái thanh toán.",
    414: "Bên cạnh các kết quả đạt được, hệ thống vẫn còn một số điểm có thể tiếp tục hoàn thiện. Chức năng vận chuyển hiện tính phí theo khu vực cố định, chưa tích hợp trực tiếp với các đơn vị giao hàng để tạo vận đơn hoặc theo dõi trạng thái giao hàng thời gian thực.",
    417: "Chức năng thanh toán QR chuyển khoản vẫn cần xác nhận thủ công, chưa có đối soát tự động với ngân hàng hoặc cổng thanh toán. Chức năng báo cáo, quản lý kho và phân quyền có thể tiếp tục mở rộng theo hướng chi tiết hơn.",
    435: "Trong tương lai, hệ thống có thể tích hợp API vận chuyển như GHN, GHTK hoặc Viettel Post để tự động tính phí vận chuyển, tạo vận đơn và theo dõi trạng thái giao hàng.",
    439: "Nhìn chung, hệ thống TienProSport đã đáp ứng phần lớn mục tiêu đề ra, hỗ trợ hoạt động bán hàng trực tuyến cho sản phẩm thể thao và phân chia trách nhiệm rõ ràng giữa khách hàng, nhân viên và quản trị viên. Mặc dù vẫn còn một số tồn đọng, hệ thống có nền tảng tốt để tiếp tục mở rộng và triển khai thực tế.",
}

SCOPE_STAFF_TEXT = (
    "Giới hạn tính năng đối với Nhân viên: Nhân viên được phép đăng nhập vào hệ thống quản trị/vận hành để thực hiện các nghiệp vụ như quản lý sản phẩm, danh mục, thuộc tính/biến thể, xử lý đơn hàng, cập nhật trạng thái đơn hàng, xác nhận thanh toán chuyển khoản nếu được phân quyền, quản lý tồn kho, nhập kho, xem lịch sử nhập kho, xem lịch sử biến động kho, theo dõi cảnh báo gần hết hàng, xử lý yêu cầu trả hàng và quản lý đánh giá sản phẩm ở mức giới hạn. Nhân viên không có quyền quản lý tài khoản người dùng cấp cao, quản lý nhân viên, phân quyền hoặc thay đổi các cấu hình hệ thống quan trọng."
)

BUSINESS_SECTION = [
    ("Các nghiệp vụ phía Khách hàng", 1),
    ("Khách hàng là người mua hàng trên website. Nhóm nghiệp vụ này tập trung vào trải nghiệm tìm kiếm, lựa chọn, đặt mua và theo dõi đơn hàng.", 0),
    ("Đăng ký, đăng nhập, đăng xuất và quản lý hồ sơ cá nhân.", 2),
    ("Xem danh sách sản phẩm, tìm kiếm, lọc sản phẩm và xem chi tiết sản phẩm.", 2),
    ("Thêm sản phẩm vào giỏ hàng, cập nhật số lượng hoặc xóa sản phẩm khỏi giỏ hàng.", 2),
    ("Đặt hàng, chọn phương thức thanh toán COD hoặc QR chuyển khoản thủ công và theo dõi trạng thái đơn hàng.", 2),
    ("Xem lịch sử đơn hàng, xem chi tiết đơn hàng, đánh giá sản phẩm sau khi nhận hàng và gửi yêu cầu trả hàng khi đủ điều kiện.", 2),
    ("Các nghiệp vụ phía Nhân viên", 1),
    ("Nhân viên là người vận hành shop, có quyền thấp hơn quản trị viên và chỉ truy cập các chức năng phục vụ xử lý nghiệp vụ hằng ngày.", 0),
    ("Đăng nhập hệ thống quản trị/vận hành và truy cập các màn hình được phân quyền.", 2),
    ("Quản lý sản phẩm, danh mục, thuộc tính và biến thể sản phẩm.", 2),
    ("Quản lý đơn hàng, cập nhật trạng thái đơn hàng và xác nhận thanh toán chuyển khoản trong phạm vi được phân quyền.", 2),
    ("Quản lý tồn kho, nhập kho, xem lịch sử nhập kho, xem lịch sử biến động kho và theo dõi cảnh báo gần hết hàng.", 2),
    ("Xử lý yêu cầu trả hàng và quản lý đánh giá sản phẩm ở mức giới hạn.", 2),
    ("Không được quản lý nhân viên, phân quyền, cấu hình hệ thống quan trọng hoặc truy cập báo cáo đầy đủ dành riêng cho quản trị viên.", 2),
    ("Các nghiệp vụ phía Quản trị viên", 1),
    ("Quản trị viên là người có quyền cao nhất, chịu trách nhiệm quản lý toàn bộ hệ thống và cấu hình vận hành.", 0),
    ("Quản lý nhân viên, khách hàng, phân quyền và các chức năng quản trị cấp cao.", 2),
    ("Quản lý cấu hình thanh toán - vận chuyển, thông tin ngân hàng/QR, thông báo và chương trình khuyến mãi.", 2),
    ("Xem dashboard, báo cáo/thống kê doanh thu, đơn hàng và các chỉ số tổng quan.", 2),
    ("Có thể thực hiện các chức năng của nhân viên khi cần để hỗ trợ vận hành hệ thống.", 2),
]

PERMISSION_ROWS = [
    ("1", "Xem sản phẩm", "Có", "Có", "Có", "Khách hàng xem trên website; nhân viên/quản trị viên xem trong màn hình quản lý."),
    ("2", "Đặt hàng", "Có", "Không", "Không", "Chỉ khách hàng thực hiện quy trình mua hàng."),
    ("3", "Xem lịch sử đơn hàng cá nhân", "Có", "Không", "Không", "Khách hàng chỉ xem đơn hàng của chính mình."),
    ("4", "Quản lý sản phẩm", "Không", "Có", "Có", "Nhân viên và quản trị viên được thêm, sửa, cập nhật sản phẩm."),
    ("5", "Quản lý danh mục", "Không", "Có", "Có", "Thuộc nhóm nghiệp vụ vận hành."),
    ("6", "Quản lý thuộc tính/biến thể", "Không", "Có", "Có", "Phục vụ quản lý kích cỡ, màu sắc và biến thể sản phẩm."),
    ("7", "Quản lý đơn hàng", "Không", "Có", "Có", "Nhân viên/quản trị viên xử lý trạng thái đơn hàng."),
    ("8", "Xác nhận thanh toán QR", "Không", "Có", "Có", "QR chuyển khoản được xác nhận thủ công trong hệ thống."),
    ("9", "Quản lý tồn kho, nhập kho", "Không", "Có", "Có", "Bao gồm tồn kho, lịch sử nhập kho và biến động kho."),
    ("10", "Xử lý trả hàng", "Có", "Có", "Có", "Khách hàng gửi yêu cầu; nhân viên/quản trị viên xử lý."),
    ("11", "Quản lý đánh giá", "Có", "Giới hạn", "Có", "Khách hàng tạo đánh giá; nhân viên kiểm duyệt/xử lý ở mức giới hạn; quản trị viên có quyền cao hơn."),
    ("12", "Quản lý khuyến mãi", "Không", "Không", "Có", "Chức năng cấp quản trị."),
    ("13", "Quản lý thông báo", "Không", "Không", "Có", "Chỉ quản trị viên cấu hình/thêm thông báo."),
    ("14", "Quản lý khách hàng", "Không", "Không", "Có", "Chỉ quản trị viên truy cập danh sách khách hàng."),
    ("15", "Quản lý nhân viên/phân quyền", "Không", "Không", "Có", "Nhân viên không được tạo tài khoản quản trị hoặc phân quyền."),
    ("16", "Cấu hình thanh toán - vận chuyển", "Không", "Không", "Có", "Bao gồm phí vận chuyển theo khu vực và thông tin QR/ngân hàng."),
    ("17", "Xem báo cáo đầy đủ", "Không", "Không", "Có", "Dashboard/báo cáo tổng quan dành cho quản trị viên."),
]

STAFF_TEST_ROWS = [
    ("TC_STAFF_01", "Nhân viên đăng nhập hệ thống quản trị/vận hành", "Tài khoản có role staff", "Đăng nhập bằng tài khoản nhân viên", "Đăng nhập thành công và được chuyển đến trang vận hành phù hợp", "Đạt"),
    ("TC_STAFF_02", "Nhân viên truy cập chức năng vận hành", "Nhân viên đã đăng nhập", "Mở các trang sản phẩm, danh mục, đơn hàng, tồn kho", "Hệ thống cho phép truy cập các chức năng được phân quyền", "Đạt"),
    ("TC_STAFF_03", "Nhân viên xử lý đơn hàng", "Có đơn hàng trong hệ thống", "Cập nhật trạng thái đơn hàng hoặc trạng thái thanh toán QR", "Hệ thống ghi nhận thay đổi trạng thái hợp lệ", "Đạt"),
    ("TC_STAFF_04", "Nhân viên quản lý kho", "Có sản phẩm/biến thể trong hệ thống", "Thực hiện nhập kho, xem lịch sử nhập kho và biến động kho", "Tồn kho và lịch sử kho được cập nhật", "Đạt"),
    ("TC_STAFF_05", "Nhân viên xử lý trả hàng/đánh giá", "Có yêu cầu trả hàng hoặc đánh giá sản phẩm", "Mở màn hình trả hàng/đánh giá và thực hiện thao tác được phân quyền", "Hệ thống cho phép xử lý ở phạm vi nhân viên", "Đạt"),
    ("TC_STAFF_06", "Nhân viên truy cập chức năng cấp quản trị", "Nhân viên đã đăng nhập", "Truy cập quản lý nhân viên, khách hàng, cấu hình hoặc báo cáo đầy đủ", "Hệ thống từ chối hoặc điều hướng về trang được phép", "Đạt"),
]


def set_run_font(run, size=BODY_SIZE, bold=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag.endswith("}r"):
            paragraph._p.remove(child)


def set_paragraph_text(paragraph, text, bold=None):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)


def style_cell(cell, size=BODY_SIZE, bold=False, align=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        if align is not None:
            paragraph.alignment = align
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold)


def set_cell_text(cell, text, size=BODY_SIZE, bold=False, align=None):
    cell.text = text
    style_cell(cell, size=size, bold=bold, align=align)


def add_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def insert_paragraph_after(anchor, text="", style=None, bold=None, align=None):
    new_p = deepcopy(anchor._p)
    anchor._p.addnext(new_p)
    paragraph = anchor._parent.paragraphs[[p._p for p in anchor._parent.paragraphs].index(new_p)]
    if style is not None:
        paragraph.style = style
    set_paragraph_text(paragraph, text, bold=bold)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def insert_new_paragraph_after(anchor, text="", style=None, bold=None, align=None):
    p = anchor._parent.add_paragraph(style=style or anchor.style)
    set_paragraph_text(p, text, bold=bold)
    if align is not None:
        p.alignment = align
    anchor._p.addnext(p._p)
    return p


def insert_table_after_paragraph(paragraph, rows, cols, style="Table Grid"):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(6.2))
    try:
        table.style = style
    except KeyError:
        pass
    paragraph._p.addnext(table._tbl)
    return table


def remove_block_between(doc, start_idx, end_idx):
    for idx in range(end_idx - 1, start_idx, -1):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)


def update_targeted_paragraphs(doc):
    for idx, text in TARGETED_PARAGRAPH_TEXT.items():
        if idx < len(doc.paragraphs):
            set_paragraph_text(doc.paragraphs[idx], text)


def add_staff_scope(doc):
    if any(SCOPE_STAFF_TEXT[:45] in p.text for p in doc.paragraphs):
        return
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.startswith("Giới hạn tính năng đối với Khách hàng"):
            insert_new_paragraph_after(paragraph, SCOPE_STAFF_TEXT, style=paragraph.style)
            return
    if len(doc.paragraphs) > 68:
        insert_new_paragraph_after(doc.paragraphs[68], SCOPE_STAFF_TEXT, style=doc.paragraphs[68].style)


def insert_business_section(doc):
    start = next((i for i, p in enumerate(doc.paragraphs) if "Hệ thống Website có các quy trình nghiệp vụ" in p.text), None)
    end = next((i for i, p in enumerate(doc.paragraphs[start + 1 :], start + 1) if "Sơ đồ chức năng" in p.text), None) if start is not None else None
    if start is None or end is None:
        return
    set_paragraph_text(doc.paragraphs[start], TARGETED_PARAGRAPH_TEXT[127])
    remove_block_between(doc, start, end)
    anchor = doc.paragraphs[start]
    for text, kind in BUSINESS_SECTION:
        if kind == 1:
            p = insert_new_paragraph_after(anchor, text, style=anchor.style, bold=True)
        elif kind == 2:
            p = insert_new_paragraph_after(anchor, "- " + text, style=anchor.style)
        else:
            p = insert_new_paragraph_after(anchor, text, style=anchor.style)
        anchor = p


def add_permission_table(doc):
    if any("Bảng phân quyền theo vai trò" in p.text for p in doc.paragraphs):
        return
    idx = next((i for i, p in enumerate(doc.paragraphs) if "Sơ đồ Use case tổng quát" in p.text), None)
    if idx is None:
        return
    anchor = doc.paragraphs[idx]
    note = insert_new_paragraph_after(
        anchor,
        "Ghi chú cập nhật theo code hiện tại: sơ đồ use case tổng quát cần thể hiện 3 actor gồm Khách hàng, Nhân viên và Quản trị viên. Khách hàng thực hiện nghiệp vụ mua hàng; Nhân viên thực hiện nghiệp vụ vận hành; Quản trị viên quản lý hệ thống, cấu hình, nhân viên và báo cáo.",
        style=anchor.style,
    )
    title = insert_new_paragraph_after(note, "Bảng phân quyền theo vai trò", style=anchor.style, bold=True)
    table = insert_table_after_paragraph(title, len(PERMISSION_ROWS) + 1, 6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)
    headers = ["STT", "Chức năng", "Khách hàng", "Nhân viên", "Quản trị viên", "Ghi chú"]
    widths = [0.35, 1.45, 0.65, 0.65, 0.75, 2.15]
    for col, header in enumerate(headers):
        set_cell_text(table.cell(0, col), header, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r, row_data in enumerate(PERMISSION_ROWS, 1):
        for c, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if c in (0, 2, 3, 4) else None
            set_cell_text(table.cell(r, c), value, size=7.2, align=align)
    for row in table.rows:
        for idx_width, width in enumerate(widths):
            row.cells[idx_width].width = Inches(width)


def update_khach_hang_table(doc):
    if not doc.tables:
        return
    table = doc.tables[0]
    # Description row may be represented by multiple cells depending on merge state.
    for cell in table.rows[0].cells:
        set_cell_text(cell, "Bảng khach_hang lưu thông tin tài khoản khách hàng, nhân viên và quản trị viên. Trường role dùng để phân biệt các vai trò customer, staff và admin; trường vai_tro được giữ để tương thích dữ liệu cũ.", size=12)
    has_role = any(row.cells[0].text.strip().lower() == "role" for row in table.rows)
    for row in table.rows:
        if row.cells[0].text.strip().lower() == "vai_tro":
            set_cell_text(row.cells[3], "Cờ vai trò cũ dùng để tương thích dữ liệu trước đây; tài khoản admin cũ được ánh xạ sang role admin.", size=12)
    if not has_role:
        row = table.add_row()
        values = ["role", "Varchar(20)", "", "Vai trò tài khoản: customer, staff hoặc admin. Giá trị này quyết định quyền truy cập frontend và backend."]
        for idx, value in enumerate(values):
            set_cell_text(row.cells[idx], value, size=12)
    add_table_borders(table)


def append_staff_tests(doc):
    if len(doc.tables) <= 33:
        return
    table = doc.tables[33]
    existing = "\n".join(cell.text for row in table.rows for cell in row.cells)
    if "TC_STAFF_01" in existing:
        return
    for row_data in STAFF_TEST_ROWS:
        row = table.add_row()
        for idx, value in enumerate(row_data[: len(row.cells)]):
            set_cell_text(row.cells[idx], value, size=10)
    add_table_borders(table)


def update_chapter5_goal_table(doc):
    if len(doc.tables) <= 36:
        return
    table = doc.tables[36]
    full = "\n".join(cell.text for row in table.rows for cell in row.cells)
    if "3 nhóm người dùng" in full:
        return
    for row in table.rows:
        text = " ".join(cell.text for cell in row.cells).lower()
        if "quản lý" in text and ("admin" in text or "quản trị" in text):
            if len(row.cells) >= 5:
                set_cell_text(row.cells[2], "Hệ thống đã hỗ trợ phân quyền 3 nhóm người dùng: khách hàng, nhân viên và quản trị viên. Nhân viên xử lý nghiệp vụ vận hành; quản trị viên quản lý hệ thống, nhân viên, cấu hình và báo cáo.", size=10)
                set_cell_text(row.cells[4], "Đã điều chỉnh theo code hiện tại. Chưa tích hợp API giao vận/thanh toán tự động; hệ thống dùng phí vận chuyển theo khu vực và QR chuyển khoản thủ công.", size=10)
            break
    add_table_borders(table)


def update_use_case_actor_rows(doc):
    for table in doc.tables[19:33]:
        name = ""
        for row in table.rows:
            if row.cells[0].text.strip().lower() in {"tên use case", "ten use case"}:
                name = row.cells[1].text.strip().lower()
                break
        actor = None
        if any(key in name for key in ["đăng ký", "đăng nhập", "chỉnh sửa", "tìm kiếm", "giỏ hàng", "thanh toán"]):
            actor = "Khách hàng"
        elif "yêu cầu trả hàng" in name and "xử lý" not in name:
            actor = "Khách hàng"
        elif any(key in name for key in ["xử lý yêu cầu trả hàng", "quản lý danh mục", "quản lý sản phẩm", "quản lý đơn hàng", "quản lý tồn kho"]):
            actor = "Nhân viên, Quản trị viên"
        elif any(key in name for key in ["khuyến mãi", "báo cáo"]):
            actor = "Quản trị viên"
        if actor:
            for row in table.rows:
                if row.cells[0].text.strip().lower() in {"tác nhân", "actor"}:
                    set_cell_text(row.cells[1], actor, size=12)
                    break


def main():
    if not INPUT.exists():
        raise FileNotFoundError(INPUT)
    doc = Document(INPUT)
    update_targeted_paragraphs(doc)
    add_staff_scope(doc)
    insert_business_section(doc)
    add_permission_table(doc)
    update_khach_hang_table(doc)
    append_staff_tests(doc)
    update_chapter5_goal_table(doc)
    update_use_case_actor_rows(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
