from pathlib import Path
import re
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


INPUT = Path(r"C:\Users\ACER\Downloads\LVTN_v2_ThemSoDoTuanTu.docx")
OUTPUT = Path(r"C:\Users\ACER\Downloads\LVTN_v2_ThemSoDoTuanTu_ThemSoDoHoatDong.docx")
IMG_DIR = Path(r"C:\Users\ACER\Downloads\banhangcodex\docs\activity_diagrams")
FONT = "Times New Roman"


def get_font(name, size):
    candidates = [
        Path(r"C:\Windows\Fonts") / name,
        Path(r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\times.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


FONT_TITLE = get_font("arialbd.ttf", 34)
FONT_BOX = get_font("arial.ttf", 26)
FONT_BOLD = get_font("arialbd.ttf", 26)
FONT_SMALL = get_font("arial.ttf", 22)


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def wrap_text(text, max_chars):
    lines = []
    for part in text.split("\n"):
        lines.extend(textwrap.wrap(part, width=max_chars) or [""])
    return lines


def text_size(draw, text, fnt):
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def draw_text_center(draw, cx, cy, text, fnt, max_chars=32, fill=(0, 0, 0)):
    lines = wrap_text(text, max_chars)
    total_h = len(lines) * 30
    y = cy - total_h / 2
    for line in lines:
        w, h = text_size(draw, line, fnt)
        draw.text((cx - w / 2, y), line, font=fnt, fill=fill)
        y += 30


def draw_round_box(draw, cx, cy, w, h, text):
    x1, y1 = cx - w / 2, cy - h / 2
    x2, y2 = cx + w / 2, cy + h / 2
    draw.rounded_rectangle((x1, y1, x2, y2), radius=18, outline=(0, 0, 0), width=3, fill=(255, 255, 255))
    draw_text_center(draw, cx, cy, text, FONT_BOX, max_chars=34)
    return (x1, y1, x2, y2)


def draw_diamond(draw, cx, cy, w, h, text):
    points = [(cx, cy - h / 2), (cx + w / 2, cy), (cx, cy + h / 2), (cx - w / 2, cy)]
    draw.polygon(points, outline=(0, 0, 0), fill=(255, 255, 255))
    draw.line(points + [points[0]], fill=(0, 0, 0), width=3)
    draw_text_center(draw, cx, cy, text, FONT_BOX, max_chars=19)
    return (cx - w / 2, cy - h / 2, cx + w / 2, cy + h / 2)


def draw_start(draw, cx, cy):
    r = 23
    draw.ellipse((cx - r, cy - r, cx + r, cy + r), fill=(0, 0, 0), outline=(0, 0, 0), width=3)
    return (cx - r, cy - r, cx + r, cy + r)


def draw_end(draw, cx, cy):
    r = 25
    draw.ellipse((cx - r, cy - r, cx + r, cy + r), outline=(0, 0, 0), width=4, fill=(255, 255, 255))
    draw.ellipse((cx - 14, cy - 14, cx + 14, cy + 14), fill=(0, 0, 0), outline=(0, 0, 0))
    return (cx - r, cy - r, cx + r, cy + r)


def draw_arrow_poly(draw, points, label=None, label_pos=None):
    for idx in range(len(points) - 1):
        draw.line((points[idx][0], points[idx][1], points[idx + 1][0], points[idx + 1][1]), fill=(0, 0, 0), width=3)
    x1, y1 = points[-2]
    x2, y2 = points[-1]
    if abs(x2 - x1) > abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        head = [(x2, y2), (x2 - direction * 16, y2 - 8), (x2 - direction * 16, y2 + 8)]
    else:
        direction = 1 if y2 > y1 else -1
        head = [(x2, y2), (x2 - 8, y2 - direction * 16), (x2 + 8, y2 - direction * 16)]
    draw.polygon(head, fill=(0, 0, 0))
    if label:
        lx, ly = label_pos if label_pos else points[len(points) // 2]
        w, h = text_size(draw, label, FONT_SMALL)
        draw.rectangle((lx - w / 2 - 6, ly - h / 2 - 4, lx + w / 2 + 6, ly + h / 2 + 4), fill=(255, 255, 255))
        draw.text((lx - w / 2, ly - h / 2), label, font=FONT_SMALL, fill=(0, 0, 0))


def draw_activity(path, title, nodes):
    width = 1800
    center_x = 760
    branch_x = 1370
    step_gap = 132
    top = 105
    bottom_pad = 100
    height = top + len(nodes) * step_gap + bottom_pad

    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw.text((45, 28), title, font=FONT_TITLE, fill=(0, 0, 0))

    positions = []
    for i, node in enumerate(nodes):
        y = top + i * step_gap
        kind = node[0]
        text = node[1] if len(node) > 1 else ""
        if kind == "start":
            bbox = draw_start(draw, center_x, y)
        elif kind == "end":
            bbox = draw_end(draw, center_x, y)
        elif kind == "decision":
            bbox = draw_diamond(draw, center_x, y, 270, 122, text)
        else:
            bbox = draw_round_box(draw, center_x, y, 610, 78, text)
        positions.append((kind, text, y, bbox, node))

    # Main vertical flow.
    for i in range(len(positions) - 1):
        kind, _, y, bbox, node = positions[i]
        next_kind, _, next_y, next_bbox, _ = positions[i + 1]
        start_y = bbox[3]
        end_y = next_bbox[1]
        label = None
        if kind == "decision":
            label = node[2].get("yes", "Hợp lệ")
        draw_arrow_poly(draw, [(center_x, start_y), (center_x, end_y)], label=label, label_pos=(center_x + 58, (start_y + end_y) / 2))

    # Orthogonal error branches to the right. No diagonal lines.
    for kind, text, y, bbox, node in positions:
        if kind != "decision" or len(node) < 3:
            continue
        meta = node[2]
        if "no" not in meta:
            continue
        error_text = meta["no"]
        box_bbox = draw_round_box(draw, branch_x, y, 520, 72, error_text)
        draw_arrow_poly(
            draw,
            [(bbox[2], y), (box_bbox[0], y)],
            label=meta.get("no_label", "Không hợp lệ"),
            label_pos=((bbox[2] + box_bbox[0]) / 2, y - 22),
        )
        if meta.get("end", True):
            end_y = y + 96
            draw_end(draw, branch_x, end_y)
            draw_arrow_poly(draw, [(branch_x, box_bbox[3]), (branch_x, end_y - 25)])

    img.save(path, "PNG")


def a(text):
    return ("action", text)


def d(text, yes="Hợp lệ", no="Hiển thị thông báo lỗi", no_label="Không hợp lệ"):
    return ("decision", text, {"yes": yes, "no": no, "no_label": no_label})


DIAGRAMS = [
    ("đăng ký", [
        ("start", ""), a('Mở form đăng ký'), a('Nhập thông tin đăng ký'), a('Bấm nút "Đăng ký"'),
        d("Kiểm tra dữ liệu", "Hợp lệ", "Thông báo dữ liệu không hợp lệ"),
        d("Email/SĐT đã tồn tại?", "Không", "Thông báo tài khoản đã tồn tại", "Có"),
        a("Mã hóa mật khẩu"), a("Lưu tài khoản mới"), a("Hiển thị đăng ký thành công"), ("end", ""),
    ]),
    ("đăng nhập", [
        ("start", ""), a('Bấm nút "Đăng nhập"'), a("Chuyển đến trang đăng nhập"), a("Nhập thông tin đăng nhập"),
        a('Bấm nút "Đăng nhập"'), d("Kiểm tra dữ liệu", "Hợp lệ", "Hiển thị thông báo lỗi"),
        d("Tài khoản/mật khẩu đúng?", "Đúng", "Hiển thị đăng nhập thất bại", "Sai"),
        a("Tạo phiên đăng nhập/token"), a("Chuyển hướng theo vai trò"), ("end", ""),
    ]),
    ("chỉnh sửa thông tin cá nhân", [
        ("start", ""), a("Mở trang thông tin cá nhân"), a("Hiển thị dữ liệu hồ sơ"), a("Chỉnh sửa thông tin"),
        a('Bấm nút "Lưu"'), d("Kiểm tra dữ liệu", "Hợp lệ", "Thông báo dữ liệu không hợp lệ"),
        a("Cập nhật thông tin vào CSDL"), a("Hiển thị cập nhật thành công"), ("end", ""),
    ]),
    ("tìm kiếm sản phẩm", [
        ("start", ""), a("Nhập từ khóa/chọn bộ lọc"), a("Gửi điều kiện tìm kiếm"), a("Truy vấn sản phẩm"),
        d("Có sản phẩm phù hợp?", "Có", "Hiển thị không có kết quả", "Không"),
        a("Hiển thị danh sách sản phẩm"), ("end", ""),
    ]),
    ("quản lý giỏ hàng", [
        ("start", ""), a("Chọn sản phẩm/biến thể"), a("Thêm/cập nhật/xóa sản phẩm"),
        d("Sản phẩm tồn tại?", "Có", "Thông báo sản phẩm không tồn tại", "Không"),
        d("Tồn kho đủ?", "Đủ", "Thông báo tồn kho không đủ", "Không đủ"),
        a("Lưu thay đổi giỏ hàng"), a("Hiển thị giỏ hàng cập nhật"), ("end", ""),
    ]),
    ("thanh toán đơn hàng", [
        ("start", ""), a("Mở trang thanh toán"), a("Nhập địa chỉ và chọn thanh toán"), a("Tính phí vận chuyển"),
        d("Giỏ hàng hợp lệ?", "Hợp lệ", "Thông báo giỏ hàng rỗng", "Không"),
        d("Tồn kho đủ?", "Đủ", "Thông báo tồn kho không đủ", "Không đủ"),
        a("Tạo đơn hàng và chi tiết"), a("Hiển thị đặt hàng thành công/QR"), ("end", ""),
    ]),
    ("yêu cầu trả hàng / hoàn tiền", [
        ("start", ""), a("Mở chi tiết đơn hàng"), a("Chọn sản phẩm cần trả"), a("Nhập lý do/hình ảnh"),
        d("Đơn đủ điều kiện?", "Có", "Thông báo đơn chưa đủ điều kiện", "Không"),
        d("Số lượng trả hợp lệ?", "Hợp lệ", "Thông báo số lượng không hợp lệ"),
        a("Lưu yêu cầu trả hàng"), a("Hiển thị gửi yêu cầu thành công"), ("end", ""),
    ]),
    ("xử lý yêu cầu trả hàng", [
        ("start", ""), a("Mở danh sách yêu cầu trả hàng"), a("Xem chi tiết yêu cầu"),
        a("Chọn duyệt hoặc từ chối"), d("Đủ quyền xử lý?", "Có", "Thông báo không đủ quyền", "Không"),
        d("Yêu cầu còn hiệu lực?", "Có", "Thông báo yêu cầu đã xử lý", "Không"),
        a("Cập nhật trạng thái yêu cầu"), a("Cập nhật kho nếu hoàn tất trả hàng"), a("Hiển thị kết quả xử lý"), ("end", ""),
    ]),
    ("quản lý danh mục", [
        ("start", ""), a("Mở trang quản lý danh mục"), a("Hiển thị danh sách danh mục"),
        a("Thêm/sửa/ẩn danh mục"), d("Dữ liệu hợp lệ?", "Hợp lệ", "Thông báo tên danh mục rỗng"),
        d("Có thể cập nhật/xóa?", "Có", "Thông báo danh mục đang được dùng", "Không"),
        a("Lưu thay đổi danh mục"), a("Hiển thị thông báo thành công"), ("end", ""),
    ]),
    ("quản lý sản phẩm", [
        ("start", ""), a("Mở trang quản lý sản phẩm"), a("Hiển thị danh sách sản phẩm"),
        a("Thêm/sửa sản phẩm và biến thể"), d("Dữ liệu hợp lệ?", "Hợp lệ", "Thông báo dữ liệu sản phẩm lỗi"),
        d("Danh mục tồn tại?", "Có", "Thông báo danh mục không tồn tại", "Không"),
        a("Lưu sản phẩm và biến thể"), a("Hiển thị thông báo thành công"), ("end", ""),
    ]),
    ("quản lý đơn hàng", [
        ("start", ""), a("Mở trang quản lý đơn hàng"), a("Xem chi tiết đơn hàng"),
        a("Chọn trạng thái cần cập nhật"), d("Đủ quyền?", "Có", "Thông báo không đủ quyền", "Không"),
        d("Trạng thái hợp lệ?", "Hợp lệ", "Thông báo sai quy trình trạng thái"),
        a("Lưu trạng thái đơn hàng"), a("Hiển thị cập nhật thành công"), ("end", ""),
    ]),
    ("quản lý khuyến mãi", [
        ("start", ""), a("Mở trang quản lý khuyến mãi"), a("Hiển thị danh sách mã"),
        a("Thêm/sửa mã khuyến mãi"), d("Dữ liệu hợp lệ?", "Hợp lệ", "Thông báo dữ liệu khuyến mãi lỗi"),
        d("Mã khuyến mãi trùng?", "Không", "Thông báo mã đã tồn tại", "Có"),
        a("Lưu mã khuyến mãi"), a("Hiển thị thông báo thành công"), ("end", ""),
    ]),
    ("quản lý tồn kho", [
        ("start", ""), a("Mở trang quản lý tồn kho"), a("Hiển thị danh sách tồn kho"),
        a("Nhập kho/điều chỉnh kho"), d("Sản phẩm tồn tại?", "Có", "Thông báo sản phẩm không tồn tại", "Không"),
        d("Số lượng hợp lệ?", "Hợp lệ", "Thông báo số lượng không hợp lệ"),
        a("Cập nhật số lượng tồn"), a("Ghi lịch sử biến động kho"), a("Hiển thị thông báo thành công"), ("end", ""),
    ]),
    ("báo cáo thống kê", [
        ("start", ""), a("Mở trang báo cáo/thống kê"), a("Chọn khoảng thời gian/bộ lọc"),
        d("Điều kiện lọc hợp lệ?", "Hợp lệ", "Thông báo khoảng thời gian không hợp lệ"),
        a("Truy vấn dữ liệu thống kê"), d("Có dữ liệu?", "Có", "Hiển thị thông báo không có dữ liệu", "Không"),
        a("Tổng hợp doanh thu, đơn hàng, tồn kho"), a("Hiển thị biểu đồ và chỉ số"), ("end", ""),
    ]),
]


def insert_paragraph_after(anchor, text="", style=None, alignment=None, size=12, italic=False):
    p = anchor._parent.add_paragraph()
    if style is not None:
        p.style = style
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, italic=italic)
    if alignment is not None:
        p.alignment = alignment
    anchor._p.addnext(p._p)
    return p


def insert_picture_after(anchor, image_path):
    p = anchor._parent.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.25))
    anchor._p.addnext(p._p)
    return p


def main():
    if not INPUT.exists():
        raise FileNotFoundError(INPUT)
    IMG_DIR.mkdir(parents=True, exist_ok=True)

    image_paths = []
    for i, (name, nodes) in enumerate(DIAGRAMS, start=1):
        path = IMG_DIR / f"activity_{i:02d}.png"
        draw_activity(path, f"Sơ đồ hoạt động {name}", nodes)
        image_paths.append(path)

    doc = Document(INPUT)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    nums = [int(m.group(1)) for m in re.finditer(r"Hình\s+3\.(\d+)", all_text)]
    next_num = max(nums or [28]) + 1

    heading = None
    for p in doc.paragraphs:
        if p.text.strip() == "Sơ đồ hoạt động":
            heading = p
            break
    if heading is None:
        raise RuntimeError("Không tìm thấy mục Sơ đồ hoạt động.")

    anchor = heading
    intro = (
        "Các sơ đồ hoạt động dưới đây mô tả luồng xử lý chính và các nhánh ngoại lệ cơ bản. "
        "Bố cục được trình bày theo hướng từ trên xuống, ưu tiên đường nối thẳng đứng và ngang để hạn chế đường chéo."
    )
    anchor = insert_paragraph_after(anchor, intro, style="Normal", size=13)

    for i, (name, _) in enumerate(DIAGRAMS):
        anchor = insert_paragraph_after(anchor, "", style="Normal")
        anchor = insert_picture_after(anchor, image_paths[i])
        caption = insert_paragraph_after(
            anchor,
            f"Hình 3.{next_num + i}. Sơ đồ hoạt động {name}",
            style="Normal",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            size=12,
            italic=True,
        )
        anchor = caption

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
