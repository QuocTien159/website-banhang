from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from copy import deepcopy


INPUT = Path(r"C:\Users\ACER\Downloads\testLVTN.docx")
OUTPUT = Path(r"C:\Users\ACER\Downloads\testLVTN_Chuong5.docx")
FONT = "Times New Roman"
BODY_SIZE = Pt(13)
HEADING1_SIZE = Pt(24)
HEADING2_SIZE = Pt(15)


GOAL_ROWS = [
    [
        "1",
        "Hệ thống hóa kiến thức đã học trong chương trình đào tạo chuyên ngành Công nghệ thông tin, đặc biệt là các quy trình phân tích, thiết kế hệ thống phần mềm và kiến trúc hệ thống thông tin.",
        "Đề tài đã trình bày quy trình khảo sát, phân tích yêu cầu, thiết kế cơ sở dữ liệu, thiết kế mô hình xử lý, xây dựng chức năng và kiểm thử hệ thống. Các nội dung này thể hiện việc vận dụng kiến thức phân tích và thiết kế hệ thống vào một bài toán thương mại điện tử cụ thể.",
        "Đạt",
        "Nội dung báo cáo đã phản ánh tương đối đầy đủ vòng đời phát triển phần mềm từ phân tích đến kiểm thử.",
    ],
    [
        "2",
        "Nghiên cứu và ứng dụng các công nghệ, framework lập trình web hiện đại (áp dụng mô hình kiến trúc phần mềm MVC) cùng hệ quản trị cơ sở dữ liệu quan hệ (MySQL) để giải quyết bài toán tự động hóa vận hành thương mại điện tử.",
        "Hệ thống được xây dựng theo hướng tách biệt frontend và backend, backend sử dụng Laravel theo mô hình MVC, dữ liệu được tổ chức theo mô hình quan hệ với các bảng khách hàng, sản phẩm, biến thể, đơn hàng, kho hàng, đánh giá và trả hàng.",
        "Đạt",
        "Các công nghệ được lựa chọn phù hợp với bài toán. Hệ thống có cấu trúc mã nguồn rõ ràng, dễ bảo trì và mở rộng.",
    ],
    [
        "3",
        "Xây dựng thành công một hệ thống Website Thương mại điện tử B2C (Business to Consumer) hoạt động ổn định, phục vụ trực tiếp cho mô hình kinh doanh trực tuyến của cửa hàng.",
        "Hệ thống TienProSport đã hỗ trợ mô hình bán hàng trực tuyến cho khách hàng cuối, bao gồm xem sản phẩm, quản lý giỏ hàng, đặt hàng, theo dõi đơn hàng, thanh toán COD, thanh toán QR chuyển khoản thủ công và các nghiệp vụ sau bán hàng.",
        "Đạt",
        "Kết quả kiểm thử ở Chương 4 cho thấy các chức năng chính hoạt động ổn định trong phạm vi dữ liệu thử nghiệm.",
    ],
    [
        "4",
        "Cung cấp giao diện người dùng (Storefront) trực quan, cho phép khách hàng dễ dàng tra cứu, lọc sản phẩm theo các tiêu chí phức tạp, tự động tính toán phí vận chuyển chính xác và đặt hàng nhanh chóng.",
        "Website đã hỗ trợ khách hàng xem danh sách sản phẩm, tìm kiếm, lọc sản phẩm, xem chi tiết sản phẩm, chọn biến thể, thêm vào giỏ hàng, đặt hàng và tính phí vận chuyển theo khu vực. Hệ thống cũng hỗ trợ thanh toán COD và QR chuyển khoản.",
        "Đạt",
        "Phí vận chuyển hiện tính theo vùng nội thành, ngoại thành và tỉnh khác. Chức năng này đáp ứng phạm vi hiện tại nhưng chưa tích hợp trực tiếp API đơn vị vận chuyển.",
    ],
    [
        "5",
        "Xây dựng phân hệ Quản trị (Admin Panel) toàn diện, tích hợp giải pháp quản lý biến thể hàng hóa và đồng bộ hóa trạng thái đơn hàng theo thời gian thực với đối tác logistics bên thứ ba (3PL).",
        "Phân hệ quản trị đã hỗ trợ quản lý danh mục, sản phẩm, biến thể, hình ảnh, đơn hàng, xác nhận thanh toán chuyển khoản, khách hàng, mã khuyến mãi, đánh giá, yêu cầu trả hàng, nhập kho, lịch sử biến động kho, cảnh báo gần hết hàng và dashboard doanh thu.",
        "Đạt một phần",
        "Các chức năng quản trị nội bộ đã được triển khai tốt. Tuy nhiên, đồng bộ trạng thái đơn hàng theo thời gian thực với đối tác logistics bên thứ ba chưa được tích hợp đầy đủ.",
    ],
]


ISSUES = [
    "Thanh toán QR chuyển khoản đã hỗ trợ tạo nội dung chuyển khoản và mã QR, tuy nhiên việc đối soát giao dịch vẫn cần quản trị viên xác nhận thủ công. Hệ thống chưa tự động kiểm tra giao dịch ngân hàng hoặc tự động cập nhật trạng thái thanh toán từ phía ngân hàng.",
    "Phí vận chuyển đang được tính theo khu vực cố định như nội thành, ngoại thành và tỉnh khác. Cách làm này phù hợp với phạm vi thử nghiệm hiện tại nhưng chưa tích hợp trực tiếp với đơn vị vận chuyển để lấy phí theo trọng lượng, kích thước, địa chỉ cụ thể và chính sách từng thời điểm.",
    "Hệ thống chưa có theo dõi vận đơn thời gian thực. Trạng thái giao hàng hiện chủ yếu do quản trị viên cập nhật, chưa nhận dữ liệu tự động từ GHN, GHTK hoặc các đơn vị vận chuyển khác thông qua API hoặc webhook.",
    "Chức năng báo cáo và thống kê đã đáp ứng các chỉ số cơ bản như số lượng đơn hàng, trạng thái đơn và doanh thu hợp lệ. Tuy nhiên, báo cáo lợi nhuận, hiệu suất sản phẩm, xu hướng mua hàng và phân tích theo thời gian vẫn còn có thể mở rộng thêm.",
    "Chức năng quản lý kho đã có nhập kho, lịch sử biến động kho và cảnh báo gần hết hàng. Trong các phiên bản tiếp theo, phân hệ này có thể bổ sung quản lý nhà cung cấp, giá nhập, công nợ nhà cung cấp và báo cáo tồn kho nâng cao.",
    "Chức năng đánh giá sản phẩm đã hỗ trợ khách hàng đánh giá sau khi đơn hàng hoàn thành và có quy trình duyệt của quản trị viên. Tuy nhiên, hệ thống có thể tiếp tục bổ sung cơ chế chống spam, lọc đánh giá theo hình ảnh, thống kê sao chi tiết và phát hiện nội dung không phù hợp.",
    "Chức năng trả hàng đã hỗ trợ khách gửi yêu cầu và quản trị viên xử lý nhập lại kho. Tuy nhiên, quy trình hoàn tiền vẫn còn mang tính thủ công, chưa tích hợp với cổng thanh toán hoặc nghiệp vụ hoàn tiền tự động.",
    "Hệ thống chưa được kiểm thử chuyên sâu với lượng người dùng lớn, số lượng đơn hàng cao hoặc nhiều thao tác đồng thời. Do đó, cần tiếp tục kiểm thử hiệu năng, độ ổn định và khả năng mở rộng trước khi triển khai trên môi trường thực tế quy mô lớn.",
    "Một số phần có thể tiếp tục tối ưu về giao diện, trải nghiệm người dùng, bảo mật, sao lưu dữ liệu và hiệu năng truy vấn cơ sở dữ liệu. Đây là các hướng cải thiện cần thiết để hệ thống vận hành ổn định hơn trong dài hạn.",
]


FUTURE = [
    "Tích hợp cổng thanh toán tự động như VNPay, PayOS, MoMo hoặc ZaloPay để tự động xác nhận giao dịch, giảm thao tác thủ công cho quản trị viên.",
    "Tích hợp API vận chuyển như GHN, GHTK hoặc Viettel Post để tự động tính phí vận chuyển, tạo vận đơn, in mã vận đơn và theo dõi quá trình giao hàng.",
    "Phát triển hệ thống voucher, mã giảm giá và chương trình khuyến mãi linh hoạt theo từng nhóm sản phẩm, từng khách hàng hoặc từng giai đoạn bán hàng.",
    "Mở rộng báo cáo doanh thu, lợi nhuận, tồn kho và hiệu suất sản phẩm theo ngày, tháng, quý, năm; hỗ trợ biểu đồ và bộ lọc nâng cao cho quản trị viên.",
    "Bổ sung quản lý nhà cung cấp, giá nhập, lịch sử nhập hàng, công nợ nhà cung cấp và các báo cáo liên quan đến chi phí hàng hóa.",
    "Phát triển chức năng hoàn tiền tự động hoặc bán tự động cho quy trình trả hàng, giúp quản trị viên theo dõi trạng thái hoàn tiền rõ ràng hơn.",
    "Nâng cấp chức năng đánh giá sản phẩm với duyệt đánh giá, thống kê sao, lọc đánh giá có hình ảnh và hỗ trợ phản hồi chi tiết từ quản trị viên.",
    "Xây dựng hệ thống gợi ý sản phẩm dựa trên hành vi xem, tìm kiếm, thêm vào giỏ hàng và lịch sử mua hàng của khách.",
    "Bổ sung phân quyền quản trị viên chi tiết theo vai trò như quản lý sản phẩm, quản lý đơn hàng, quản lý kho, chăm sóc khách hàng và quản trị hệ thống.",
    "Tối ưu bảo mật, hiệu năng, sao lưu dữ liệu và khả năng mở rộng để hệ thống có thể vận hành ổn định khi lượng người dùng và dữ liệu tăng lên.",
    "Phát triển ứng dụng mobile hoặc PWA để khách hàng có thể mua sắm thuận tiện hơn trên thiết bị di động.",
    "Tích hợp email, SMS hoặc Zalo notification để tự động thông báo trạng thái đơn hàng, thanh toán, giao hàng, khuyến mãi và yêu cầu trả hàng cho khách.",
]

LIST_NUM_PR = None


def set_run_font(run, size=BODY_SIZE, bold=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = size
    if bold is not None:
        run.bold = bold


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph(style=f"Heading {level}")
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        size = HEADING1_SIZE
    else:
        size = HEADING2_SIZE
    clear_paragraph(p)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)
    p.paragraph_format.space_after = Pt(0)
    return p


def add_body(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Paragraph")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if LIST_NUM_PR is not None:
        p_pr = p._p.get_or_add_pPr()
        old_num_pr = p_pr.numPr
        if old_num_pr is not None:
            p_pr.remove(old_num_pr)
        p_pr.append(deepcopy(LIST_NUM_PR))
    r = p.add_run(text)
    set_run_font(r)
    return p


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, size=Pt(11)):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_cell_margins(table, top=80, start=100, bottom=80, end=100):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str = "D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_goal_table(doc: Document):
    headers = [
        "STT",
        "Mục tiêu đề ra ở mục 1.4",
        "Kết quả thực hiện trong hệ thống",
        "Đánh giá",
        "Ghi chú/Giải thích",
    ]
    widths = [0.42, 2.05, 2.05, 0.72, 1.26]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell)
        set_cell_text(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(10))
    for row in GOAL_ROWS:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, align=align, size=Pt(10))
    doc.add_paragraph()
    return table


def remove_existing_empty_chapter5(doc: Document):
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1" and p.text.strip().upper() == "KẾT LUẬN":
            start = i
    if start is None:
        return
    for p in list(doc.paragraphs[start:]):
        parent = p._element.getparent()
        parent.remove(p._element)


def build():
    global LIST_NUM_PR
    doc = Document(INPUT)
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "List Paragraph" and paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            LIST_NUM_PR = deepcopy(paragraph._p.pPr.numPr)
            break
    remove_existing_empty_chapter5(doc)

    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)

    add_heading(doc, "CHƯƠNG 5. KẾT LUẬN", 1)

    add_heading(doc, "5.1. KẾT QUẢ ĐỐI CHIẾU VỚI MỤC TIÊU", 2)
    add_body(
        doc,
        "Mục này đối chiếu kết quả xây dựng hệ thống TienProSport với các mục tiêu đã nêu ở mục 1.4 của luận văn. Việc đánh giá được thực hiện dựa trên các chức năng đã triển khai, kết quả kiểm thử ở Chương 4 và phạm vi hiện thực của hệ thống.",
    )
    add_goal_table(doc)
    add_body(
        doc,
        "Nhìn chung, hệ thống đã hoàn thành phần lớn các mục tiêu chính của đề tài. Các chức năng phục vụ trực tiếp cho hoạt động bán hàng trực tuyến như xem sản phẩm, tìm kiếm, lọc sản phẩm, quản lý giỏ hàng, đặt hàng, thanh toán COD, thanh toán QR chuyển khoản, tính phí vận chuyển theo khu vực và xem lịch sử đơn hàng đã được triển khai. Ở phía quản trị viên, hệ thống đã hỗ trợ quản lý danh mục, sản phẩm, biến thể, khách hàng, đơn hàng, thanh toán, kho hàng, đánh giá, trả hàng và thống kê doanh thu.",
    )
    add_body(
        doc,
        "Bên cạnh đó, một số mục tiêu có phạm vi rộng như đồng bộ trạng thái đơn hàng theo thời gian thực với đối tác logistics bên thứ ba mới chỉ đạt một phần. Hệ thống hiện đã có nền tảng quản lý đơn hàng, kho hàng và phí vận chuyển, nhưng chưa tích hợp API vận chuyển thực tế để tự động tạo vận đơn, lấy phí vận chuyển theo đơn vị giao hàng và cập nhật trạng thái giao hàng theo thời gian thực.",
    )
    add_body(
        doc,
        "Về mức độ đáp ứng, hệ thống phù hợp với yêu cầu xây dựng một website bán đồ thể thao có khả năng vận hành trong phạm vi cửa hàng trực tuyến. Các chức năng cốt lõi đã đáp ứng tốt yêu cầu đề tài, trong khi các chức năng tích hợp nâng cao có thể tiếp tục được hoàn thiện ở các phiên bản sau.",
    )

    add_heading(doc, "5.2. CÁC VẤN ĐỀ CÒN TỒN ĐỌNG", 2)
    add_body(
        doc,
        "Trong quá trình xây dựng và kiểm thử, hệ thống TienProSport đã đáp ứng được các nghiệp vụ chính của một website bán hàng trực tuyến. Tuy nhiên, để có thể triển khai ở quy mô thực tế lớn hơn, hệ thống vẫn còn một số vấn đề cần tiếp tục cải thiện như sau:",
    )
    for item in ISSUES:
        add_bullet(doc, item)
    add_body(
        doc,
        "Những tồn đọng trên không làm thay đổi kết quả đạt được của đề tài, mà thể hiện các giới hạn hợp lý trong phạm vi luận văn tốt nghiệp. Đây cũng là cơ sở để định hướng phát triển hệ thống trong các giai đoạn tiếp theo.",
    )

    add_heading(doc, "5.3. MỞ RỘNG / HƯỚNG PHÁT TRIỂN", 2)
    add_body(
        doc,
        "Từ nền tảng đã xây dựng, hệ thống TienProSport có thể tiếp tục được mở rộng theo nhiều hướng nhằm nâng cao mức độ tự động hóa, tối ưu trải nghiệm người dùng và tăng khả năng vận hành thực tế:",
    )
    for item in FUTURE:
        add_bullet(doc, item)
    add_body(
        doc,
        "Tóm lại, hệ thống website bán hàng TienProSport đã đáp ứng phần lớn mục tiêu đề ra trong luận văn. Hệ thống hỗ trợ các hoạt động bán hàng trực tuyến cho sản phẩm thể thao, bao gồm quản lý sản phẩm, biến thể, giỏ hàng, đơn hàng, thanh toán, vận chuyển, kho hàng và các nghiệp vụ sau bán hàng. Mặc dù vẫn còn một số tồn đọng cần tiếp tục hoàn thiện, hệ thống đã có nền tảng tốt để mở rộng thêm các chức năng tự động hóa và triển khai thực tế trong tương lai.",
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
