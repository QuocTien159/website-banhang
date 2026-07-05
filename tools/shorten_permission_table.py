from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SRC = Path(r"C:\Users\ACER\Downloads\LVTN_CapNhat_3Actor.docx")
OUT = Path(r"C:\Users\ACER\Downloads\LVTN_CapNhat_3Actor_BangPhanQuyenGon.docx")
FONT = "Times New Roman"

HEADERS = ["STT", "Nhóm chức năng", "Khách hàng", "Nhân viên", "Quản trị viên", "Ghi chú"]
ROWS = [
    ("1", "Mua hàng", "Xem sản phẩm, đặt hàng, thanh toán, xem đơn cá nhân.", "Hỗ trợ kiểm tra khi vận hành.", "Hỗ trợ kiểm tra khi cần.", "Chức năng chính của khách hàng."),
    ("2", "Sản phẩm - danh mục", "Xem sản phẩm.", "Quản lý sản phẩm, danh mục, thuộc tính/biến thể.", "Toàn quyền quản lý.", "Nhóm nghiệp vụ vận hành."),
    ("3", "Đơn hàng - thanh toán", "Tạo và theo dõi đơn hàng cá nhân.", "Xử lý đơn, cập nhật trạng thái, xác nhận QR thủ công.", "Toàn quyền xử lý/giám sát.", "QR chưa tự động đối soát ngân hàng."),
    ("4", "Kho hàng", "Không.", "Nhập kho, xem tồn kho, lịch sử kho, cảnh báo gần hết hàng.", "Toàn quyền quản lý kho.", "Không tích hợp kho bên thứ ba."),
    ("5", "Đánh giá - trả hàng", "Gửi đánh giá và yêu cầu trả hàng.", "Xử lý theo phạm vi được phân quyền.", "Toàn quyền xử lý.", "Quản trị viên có quyền cao hơn."),
    ("6", "Khuyến mãi - thông báo", "Xem thông tin hiển thị.", "Không.", "Quản lý khuyến mãi và thông báo.", "Chức năng cấp quản trị."),
    ("7", "Người dùng - nhân viên", "Quản lý hồ sơ cá nhân.", "Không.", "Quản lý khách hàng, nhân viên, phân quyền.", "Nhân viên không quản lý tài khoản cấp cao."),
    ("8", "Cấu hình - báo cáo", "Không.", "Không.", "Cấu hình thanh toán/vận chuyển, xem dashboard và báo cáo.", "Báo cáo đầy đủ dành cho quản trị viên."),
]
WIDTHS = [0.35, 1.05, 1.35, 1.45, 1.25, 1.25]


def set_run_font(run, size=9, bold=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold


def set_cell(cell, text, size=9, bold=False, align=None):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        if align is not None:
            paragraph.alignment = align
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold)


def add_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "6")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "000000")


def remove_existing_permission_tables(doc):
    for table in list(doc.tables):
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "Khách hàng" in text and "Nhân viên" in text and "Quản trị viên" in text and (
            "Quản lý khuyến mãi" in text or "Xem báo cáo đầy đủ" in text or "Nhóm chức năng" in text
        ):
            table._tbl.getparent().remove(table._tbl)


def main():
    doc = Document(SRC)
    remove_existing_permission_tables(doc)

    title_idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        if "Bảng phân quyền theo vai trò" in paragraph.text:
            title_idx = i
            break
    if title_idx is None:
        raise RuntimeError("Không tìm thấy tiêu đề bảng phân quyền.")

    note = doc.paragraphs[title_idx - 1] if title_idx > 0 else None
    if note is not None and note.text.startswith("Ghi chú"):
        note.text = (
            "Ghi chú: Sơ đồ use case tổng quát cần thể hiện 3 actor: Khách hàng, "
            "Nhân viên và Quản trị viên. Bảng dưới đây tóm tắt quyền chính của từng vai trò."
        )
        for run in note.runs:
            set_run_font(run, size=13)

    title = doc.paragraphs[title_idx]
    table = doc.add_table(rows=len(ROWS) + 1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title._p.addnext(table._tbl)
    add_borders(table)

    for col, header in enumerate(HEADERS):
        set_cell(table.cell(0, col), header, size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_idx, row_data in enumerate(ROWS, start=1):
        for col, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if col == 0 else None
            set_cell(table.cell(row_idx, col), value, size=8.2, align=align)

    for row in table.rows:
        for col, width in enumerate(WIDTHS):
            row.cells[col].width = Inches(width)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
